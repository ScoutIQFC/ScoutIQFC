import streamlit as st
import sqlite3
import pandas as pd
import plotly.graph_objects as go
import io
import os
import base64
from dotenv import load_dotenv
load_dotenv()
# ── LOGIN ──
def check_password():
    def password_entered():
        user = st.session_state.get("login_user", "")
        pw = st.session_state.get("login_pass", "")
        try:
            admin_user = st.secrets["USERNAME"]
            admin_pass = st.secrets["PASSWORD"]
        except Exception:
            admin_user = "admin"
            admin_pass = "scoutiq2024"
        # Admin account - full access
        if user == admin_user and pw == admin_pass:
            st.session_state["authenticated"] = True
            st.session_state["account_type"] = "admin"
        # Beta account - limited access
        elif user == "admin1" and pw == "scout2024beta":
            st.session_state["authenticated"] = True
            st.session_state["account_type"] = "beta"
        else:
            st.session_state["auth_error"] = True
    if st.session_state.get("authenticated"):
        return True
    if "login_tab" not in st.session_state:
        st.session_state["login_tab"] = "signin"
    if "show_forgot" not in st.session_state:
        st.session_state["show_forgot"] = False
    st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700;800&family=Playfair+Display:ital,wght@0,700;1,400&display=swap');
html,body,[class*="css"],.stApp{background:#04080f !important;font-family:'Outfit',sans-serif !important;}
#MainMenu,footer,header{visibility:hidden !important;}
.block-container{padding:0 !important;max-width:100% !important;}
.stApp{min-height:100vh;}

/* Background */
.lbg{position:fixed;top:0;left:0;right:0;bottom:0;z-index:0;
background:radial-gradient(ellipse 80% 60% at 10% 10%,rgba(20,45,120,0.5) 0%,transparent 60%),
radial-gradient(ellipse 60% 70% at 90% 90%,rgba(8,16,60,0.55) 0%,transparent 60%),#04080f;}
.lbg::before{content:'';position:absolute;inset:0;
background-image:linear-gradient(rgba(255,255,255,0.022) 1px,transparent 1px),
linear-gradient(90deg,rgba(255,255,255,0.022) 1px,transparent 1px);
background-size:58px 58px;}

/* Inputs - WHITE bg, DARK text */
.stTextInput label,.stSelectbox label,.stTextArea label{display:none !important;}
.stTextInput > div > div{background:#ffffff !important;border:1.5px solid #e2e8f0 !important;border-radius:10px !important;}
.stTextInput > div > div:focus-within{background:#ffffff !important;border-color:#f5c842 !important;box-shadow:0 0 0 3px rgba(245,200,66,0.15) !important;}
.stTextInput input{color:#0d1117 !important;font-family:'Outfit',sans-serif !important;font-size:14px !important;padding:12px 14px !important;}
.stTextInput input::placeholder{color:#94a3b8 !important;}
.stTextArea textarea{color:#0d1117 !important;background:#ffffff !important;border:1.5px solid #e2e8f0 !important;border-radius:10px !important;font-size:13px !important;}
.stSelectbox > div > div{background:#ffffff !important;border:1.5px solid #e2e8f0 !important;border-radius:10px !important;color:#0d1117 !important;}

/* ALL buttons default - pill shape, subtle */
.stButton > button{background:rgba(255,255,255,0.08) !important;color:rgba(255,255,255,0.7) !important;
border:1px solid rgba(255,255,255,0.13) !important;border-radius:100px !important;
font-family:'Outfit',sans-serif !important;font-size:12px !important;font-weight:600 !important;
padding:9px 16px !important;width:100% !important;margin-top:0 !important;
box-shadow:none !important;letter-spacing:0.2px !important;}
.stButton > button:hover{background:rgba(255,255,255,0.13) !important;color:#fff !important;transform:none !important;box-shadow:none !important;}
.stButton > button:focus,.stButton > button:active{box-shadow:none !important;outline:none !important;}

/* Gold primary CTA */
.btn-gold .stButton > button{background:linear-gradient(135deg,#f5c842,#e6a817) !important;
color:#04080f !important;border:none !important;font-weight:700 !important;
font-size:14px !important;padding:13px 0 !important;margin-top:16px !important;
box-shadow:0 4px 20px rgba(245,200,66,0.2) !important;}
.btn-gold .stButton > button:hover{box-shadow:0 6px 28px rgba(245,200,66,0.3) !important;transform:translateY(-1px) !important;}

/* Ghost subtle */
.btn-ghost .stButton > button{background:transparent !important;color:rgba(255,255,255,0.35) !important;
border:1px solid rgba(255,255,255,0.08) !important;font-size:11px !important;padding:8px 0 !important;margin-top:6px !important;}
.btn-ghost .stButton > button:hover{color:rgba(255,255,255,0.65) !important;background:rgba(255,255,255,0.04) !important;}

/* Divider row buttons */
.btn-row .stButton > button{font-size:12px !important;padding:9px 12px !important;margin-top:0 !important;}

/* Error/success */
.epill{background:rgba(239,68,68,0.1);border:1px solid rgba(239,68,68,0.25);border-radius:8px;
padding:9px 14px;font-size:12px;color:#fca5a5;text-align:center;margin-bottom:12px;}
.spill{background:rgba(22,163,74,0.1);border:1px solid rgba(22,163,74,0.25);border-radius:8px;
padding:9px 14px;font-size:12px;color:#86efac;text-align:center;margin-bottom:12px;}

/* Stats */
.lstats{display:flex;justify-content:center;gap:32px;margin-top:24px;padding-top:20px;border-top:1px solid rgba(255,255,255,0.05);}
.lstat{text-align:center;}
.lstat-n{font-size:17px;font-weight:700;color:#f5c842;line-height:1;}
.lstat-l{font-size:9px;color:rgba(255,255,255,0.2);letter-spacing:1.5px;text-transform:uppercase;margin-top:3px;}
</style>
<div class="lbg"></div>
    """, unsafe_allow_html=True)
    tab = st.session_state.get("login_tab", "signin")
    # REQUEST DEMO PAGE
    if tab in ("demo", "demo_sent"):
        _, col, _ = st.columns([1, 1.4, 1])
        with col:
            st.markdown('<div style="position:relative;z-index:2;padding-top:48px;">', unsafe_allow_html=True)
            if tab == "demo_sent":
                st.markdown("""
                <div style="text-align:center;padding:60px 0;">
                    <div style="font-size:48px;margin-bottom:16px;">✅</div>
                    <div style="font-size:20px;font-weight:700;color:#fff;margin-bottom:8px;">Request Received</div>
                    <div style="font-size:13px;color:rgba(255,255,255,0.4);line-height:1.7;margin-bottom:24px;">
                        Our team will be in touch within 24 hours.
                    </div>
                </div>
                """, unsafe_allow_html=True)
                st.markdown('<div class="btn-ghost">', unsafe_allow_html=True)
                if st.button("Back to Sign In", key="bk_ds"):
                    st.session_state["login_tab"] = "signin"; st.rerun()
                st.markdown('</div>', unsafe_allow_html=True)
            else:
                st.markdown("""
                <div style="text-align:center;margin-bottom:28px;">
                    <div style="font-family:'Playfair Display',serif;font-size:36px;color:#fff;font-weight:700;letter-spacing:-0.5px;">
                        Scout IQ·<span style="color:#f5c842;font-style:italic;">FC</span>
                    </div>
                    <div style="font-size:16px;font-weight:600;color:rgba(255,255,255,0.7);margin-top:12px;">Request a Demo</div>
                    <div style="font-size:12px;color:rgba(255,255,255,0.35);margin-top:4px;">Get a personalised walkthrough within 24 hours</div>
                </div>
                """, unsafe_allow_html=True)
                st.markdown('<div style="font-size:10px;font-weight:700;color:rgba(255,255,255,0.35);letter-spacing:2px;text-transform:uppercase;margin-bottom:4px;">Full Name</div>', unsafe_allow_html=True)
                d_name = st.text_input("dn", placeholder="Your full name", label_visibility="collapsed", key="demo_name")
                st.markdown('<div style="font-size:10px;font-weight:700;color:rgba(255,255,255,0.35);letter-spacing:2px;text-transform:uppercase;margin-bottom:4px;margin-top:12px;">Email Address</div>', unsafe_allow_html=True)
                d_email = st.text_input("de", placeholder="your@email.com", label_visibility="collapsed", key="demo_email")
                st.markdown('<div style="font-size:10px;font-weight:700;color:rgba(255,255,255,0.35);letter-spacing:2px;text-transform:uppercase;margin-bottom:4px;margin-top:12px;">Academy / Club</div>', unsafe_allow_html=True)
                d_club = st.text_input("dc", placeholder="Club or academy name", label_visibility="collapsed", key="demo_club")
                st.markdown('<div style="font-size:10px;font-weight:700;color:rgba(255,255,255,0.35);letter-spacing:2px;text-transform:uppercase;margin-bottom:4px;margin-top:12px;">Message</div>', unsafe_allow_html=True)
                d_msg = st.text_area("dm", placeholder="Tell us how we can help...", label_visibility="collapsed", key="demo_msg", height=80)
                st.markdown('<div class="btn-gold">', unsafe_allow_html=True)
                if st.button("Send Request", key="submit_demo"):
                    if d_name and d_email:
                        st.session_state["login_tab"] = "demo_sent"; st.rerun()
                    else:
                        st.markdown('<div class="epill">Please fill in name and email</div>', unsafe_allow_html=True)
                st.markdown('</div><div class="btn-ghost">', unsafe_allow_html=True)
                if st.button("Back to Sign In", key="bk_demo"):
                    st.session_state["login_tab"] = "signin"; st.rerun()
                st.markdown('</div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)
        return False
    # CREATE ACCOUNT PAGE
    if tab == "create":
        _, col, _ = st.columns([1, 1.4, 1])
        with col:
            st.markdown('<div style="position:relative;z-index:2;padding-top:80px;text-align:center;">', unsafe_allow_html=True)
            st.markdown("""
            <div style="font-size:48px;margin-bottom:16px;">🚀</div>
            <div style="font-size:20px;font-weight:700;color:#fff;margin-bottom:8px;">Account Creation Coming Soon</div>
            <div style="font-size:13px;color:rgba(255,255,255,0.4);line-height:1.7;margin-bottom:28px;">
                Request a demo to get early access to Scout IQ·FC.
            </div>
            """, unsafe_allow_html=True)
            st.markdown('<div class="btn-gold">', unsafe_allow_html=True)
            if st.button("Request a Demo", key="goto_demo"):
                st.session_state["login_tab"] = "demo"; st.rerun()
            st.markdown('</div><div class="btn-ghost">', unsafe_allow_html=True)
            if st.button("Back to Sign In", key="bk_cr"):
                st.session_state["login_tab"] = "signin"; st.rerun()
            st.markdown('</div></div>', unsafe_allow_html=True)
        return False
    # MAIN SIGN IN - everything on one screen, no scroll
    _, col, _ = st.columns([1, 1.4, 1])
    with col:
        st.markdown('<div style="position:relative;z-index:2;">', unsafe_allow_html=True)
        # Brand — compact
        st.markdown("""
        <div style="text-align:center;padding-top:36px;margin-bottom:28px;">
            <div style="font-size:10px;font-weight:600;letter-spacing:4px;color:rgba(245,200,66,0.6);text-transform:uppercase;margin-bottom:8px;">Talent Intelligence Platform</div>
            <div style="font-family:'Playfair Display',serif;font-size:48px;font-weight:700;color:#fff;line-height:1;letter-spacing:-1px;margin-bottom:8px;">
                Scout IQ·<span style="color:#f5c842;font-style:italic;">FC</span>
            </div>
            <div style="font-size:11px;font-weight:500;color:rgba(255,255,255,0.38);letter-spacing:4px;text-transform:uppercase;">Where Talent Is Unearthed</div>
        </div>
        """, unsafe_allow_html=True)
        # Login form - no card, just the fields
        if st.session_state.get("auth_error"):
            st.markdown('<div class="epill">Incorrect username or password</div>', unsafe_allow_html=True)
        if st.session_state.get("show_forgot"):
            st.markdown('<div class="spill">Password reset instructions sent to your email</div>', unsafe_allow_html=True)
            st.session_state["show_forgot"] = False
        # Title - same style as field labels but slightly larger
        st.markdown('<div style="font-size:11px;font-weight:700;color:rgba(255,255,255,0.38);letter-spacing:3px;text-transform:uppercase;margin-bottom:16px;">Sign in to Scout IQ·FC</div>', unsafe_allow_html=True)
        # Username
        st.markdown('<div style="font-size:10px;font-weight:700;color:rgba(255,255,255,0.38);letter-spacing:2px;text-transform:uppercase;margin-bottom:4px;">Username</div>', unsafe_allow_html=True)
        st.text_input("u", key="login_user", placeholder="Enter your username", label_visibility="collapsed")
        # Password
        st.markdown('<div style="font-size:10px;font-weight:700;color:rgba(255,255,255,0.38);letter-spacing:2px;text-transform:uppercase;margin-bottom:4px;margin-top:12px;">Password</div>', unsafe_allow_html=True)
        st.text_input("p", type="password", key="login_pass", placeholder="Enter your password", label_visibility="collapsed")
        # Sign in button - full width gold
        st.markdown('<div class="btn-gold">', unsafe_allow_html=True)
        st.button("Sign In →", on_click=password_entered)
        st.markdown('</div>', unsafe_allow_html=True)
        # Divider
        st.markdown('<div style="border-top:1px solid rgba(255,255,255,0.07);margin:14px 0 12px;"></div>', unsafe_allow_html=True)
        # Three buttons in one row: Forgot Password | Create Account | Request Demo
        b1, b2, b3 = st.columns(3)
        with b1:
            st.markdown('<div class="btn-row">', unsafe_allow_html=True)
            if st.button("Forgot password?", key="forgot"):
                st.session_state["show_forgot"] = True; st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)
        with b2:
            st.markdown('<div class="btn-row">', unsafe_allow_html=True)
            if st.button("Create Account", key="nav_create"):
                st.session_state["login_tab"] = "create"; st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)
        with b3:
            st.markdown('<div class="btn-row">', unsafe_allow_html=True)
            if st.button("Request Demo", key="nav_demo"):
                st.session_state["login_tab"] = "demo"; st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)
        # Stats strip
        st.markdown("""
        <div class="lstats">
            <div class="lstat"><div class="lstat-n">100+</div><div class="lstat-l">Players</div></div>
            <div class="lstat"><div class="lstat-n">AI</div><div class="lstat-l">Reports</div></div>
            <div class="lstat"><div class="lstat-n">3+</div><div class="lstat-l">Academies</div></div>
            <div class="lstat"><div class="lstat-n">Beta</div><div class="lstat-l">Live</div></div>
        </div>
        """, unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
    return False
if not check_password():
    st.stop()
# ── DATABASE ──
conn = sqlite3.connect("scout_agent.db", check_same_thread=False)
cursor = conn.cursor()
cursor.executescript("""
CREATE TABLE IF NOT EXISTS players (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL, date_of_birth TEXT, age_group TEXT,
    position TEXT, dominant_foot TEXT, club TEXT, nationality TEXT
);
CREATE TABLE IF NOT EXISTS sessions (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    player_id INTEGER REFERENCES players(id),
    session_date TEXT, session_type TEXT, minutes_played INTEGER,
    distance_covered_km REAL, sprint_count INTEGER, top_speed_kmh REAL,
    passes_completed INTEGER, passes_attempted INTEGER, dribbles_completed INTEGER,
    defensive_actions INTEGER, goals INTEGER, assists INTEGER,
    chances_created INTEGER, tackles_won INTEGER, coachability_rating INTEGER,
    attitude_score INTEGER, consistency_rating INTEGER, coach_notes TEXT
);
CREATE TABLE IF NOT EXISTS reports (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    player_id INTEGER REFERENCES players(id),
    report_text TEXT, created_at TEXT DEFAULT CURRENT_TIMESTAMP
);
CREATE TABLE IF NOT EXISTS report_edits (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    player_id INTEGER, edited_report TEXT, coach_notes TEXT,
    edited_at TEXT DEFAULT CURRENT_TIMESTAMP
);
CREATE TABLE IF NOT EXISTS epl_players (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT NOT NULL, team TEXT, position TEXT, nationality TEXT, age INTEGER
);
CREATE TABLE IF NOT EXISTS epl_sessions (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    player_id INTEGER, gameweek INTEGER, opponent TEXT, venue TEXT,
    minutes_played INTEGER DEFAULT 0, goals INTEGER DEFAULT 0,
    assists INTEGER DEFAULT 0, shots INTEGER DEFAULT 0,
    shots_on_target INTEGER DEFAULT 0, xg REAL DEFAULT 0, xa REAL DEFAULT 0,
    passes_completed INTEGER DEFAULT 0, passes_attempted INTEGER DEFAULT 0,
    key_passes INTEGER DEFAULT 0, progressive_passes INTEGER DEFAULT 0,
    dribbles_completed INTEGER DEFAULT 0, dribbles_attempted INTEGER DEFAULT 0,
    touches INTEGER DEFAULT 0, ball_recoveries INTEGER DEFAULT 0,
    tackles_won INTEGER DEFAULT 0, interceptions INTEGER DEFAULT 0,
    clearances INTEGER DEFAULT 0, yellow_cards INTEGER DEFAULT 0,
    red_cards INTEGER DEFAULT 0, rating REAL DEFAULT 0
);
CREATE TABLE IF NOT EXISTS epl_reports (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    player_id INTEGER, report_text TEXT,
    created_at TEXT DEFAULT CURRENT_TIMESTAMP
);
""")
conn.commit()
try:
    from demo_data import seed_demo_data
    seed_demo_data(conn)
except Exception:
    pass
st.set_page_config(
    page_title="Scout IQ·FC",
    page_icon="⚽",
    layout="wide",
    initial_sidebar_state="expanded"
)
# Account state
for k, v in {
    "show_account": False,
    "account_page": "overview",
    "show_edit_academy": False,
    "edit_academy_name": None,
    "show_edit_player": False,
    "edit_player_id": None,
    "confirm_del_academy": None,
    "confirm_del_player": None,
}.items():
    if k not in st.session_state:
        st.session_state[k] = v
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700;800&family=Playfair+Display:ital,wght@0,600;0,700;1,400&display=swap');

/* ── RESET & BASE ── */
*, *::before, *::after { box-sizing: border-box; }
* { font-family: 'Outfit', sans-serif !important; }
html, body, [class*="css"] { background: #f8f9fc !important; color: #0d1117; }
#MainMenu, footer, header { visibility: hidden; }
.stApp { background: #f8f9fc !important; }
.block-container { padding: 0 2rem 2rem 2rem !important; max-width: 100% !important; }

/* ── SIDEBAR ── */
[data-testid="stSidebarCollapseButton"] { display: none !important; }
[data-testid="collapsedControl"] { display: none !important; }
[data-testid="stSidebar"] {
    transform: none !important;
    display: flex !important;
    visibility: visible !important;
    min-width: 260px !important;
    max-width: 260px !important;
    background: #04080f !important;
    box-shadow: 1px 0 0 rgba(255,255,255,0.04), 4px 0 32px rgba(0,0,0,0.4) !important;
}
[data-testid="stSidebar"] > div { padding: 0 !important; }

/* Sidebar logo */
.sb-logo {
    padding: 24px 20px 20px;
    border-bottom: 1px solid rgba(255,255,255,0.06);
}
.sb-logo-name {
    font-family: 'Playfair Display', serif !important;
    font-size: 20px; font-weight: 700; color: #ffffff;
    letter-spacing: -0.3px; line-height: 1;
}
.sb-logo-name .fc { color: #f5c842; font-style: italic; }
.sb-logo-tag {
    font-size: 8px; font-weight: 500;
    color: rgba(255,255,255,0.3);
    letter-spacing: 3px; text-transform: uppercase;
    margin-top: 5px;
}

/* Sidebar sections */
.sb-section {
    padding: 16px 20px 4px;
    font-size: 9px; font-weight: 700;
    color: rgba(255,255,255,0.2);
    letter-spacing: 3px; text-transform: uppercase;
}
.sb-divider { border: none; border-top: 1px solid rgba(255,255,255,0.05); margin: 6px 0; }

/* Sidebar scrollable area */
[data-testid="stSidebar"] [data-testid="stVerticalBlock"] {
    overflow-y: auto !important;
    max-height: calc(100vh - 180px) !important;
}

/* Academy row hover effects */
.academy-row { position: relative; }
.academy-row:hover .academy-actions { opacity: 1 !important; }
.academy-actions { opacity: 0; transition: opacity 0.2s; }

/* Player row hover */
.player-hover-row { position: relative; }
.player-hover-row:hover .player-actions { opacity: 1 !important; }
.player-actions { opacity: 0; transition: opacity 0.2s; }

/* Action buttons in sidebar */
.sb-action-del .stButton > button {
    background: rgba(220,38,38,0.15) !important;
    color: rgba(255,100,100,0.8) !important;
    border: 1px solid rgba(220,38,38,0.2) !important;
    border-radius: 4px !important;
    font-size: 9px !important; font-weight: 700 !important;
    padding: 3px 6px !important;
    margin: 0 !important; width: auto !important;
    box-shadow: none !important;
}
.sb-action-del .stButton > button:hover {
    background: rgba(220,38,38,0.3) !important;
    color: #ff8080 !important;
}
.sb-action-edit .stButton > button {
    background: rgba(245,200,66,0.1) !important;
    color: rgba(245,200,66,0.7) !important;
    border: 1px solid rgba(245,200,66,0.2) !important;
    border-radius: 4px !important;
    font-size: 9px !important; font-weight: 700 !important;
    padding: 3px 6px !important;
    margin: 0 !important; width: auto !important;
    box-shadow: none !important;
}
.sb-action-edit .stButton > button:hover {
    background: rgba(245,200,66,0.2) !important;
    color: #f5c842 !important;
}

/* Sidebar inputs */
[data-testid="stSidebar"] label { display: none !important; }
[data-testid="stSidebar"] .stTextInput { padding: 6px 16px; }
[data-testid="stSidebar"] .stTextInput > div > div {
    background: rgba(255,255,255,0.06) !important;
    border: 1px solid rgba(255,255,255,0.1) !important;
    border-radius: 8px !important;
}
[data-testid="stSidebar"] .stTextInput input {
    color: #fff !important; font-size: 12px !important;
    padding: 8px 12px !important;
}
[data-testid="stSidebar"] .stTextInput input::placeholder {
    color: rgba(255,255,255,0.25) !important;
}

/* Sidebar buttons */
[data-testid="stSidebar"] .stButton { margin: 0 !important; padding: 0 !important; }
[data-testid="stSidebar"] .stButton > button {
    background: transparent !important;
    color: rgba(255,255,255,0.5) !important;
    border: none !important; outline: none !important;
    box-shadow: none !important;
    font-size: 13px !important; font-weight: 400 !important;
    padding: 8px 20px !important; border-radius: 0 !important;
    width: 100% !important; text-align: left !important;
    border-left: 2px solid transparent !important;
    transition: all 0.15s !important;
    white-space: nowrap !important; overflow: hidden !important;
    text-overflow: ellipsis !important;
}
[data-testid="stSidebar"] .stButton > button:hover {
    background: rgba(255,255,255,0.05) !important;
    color: #ffffff !important;
    border-left-color: rgba(245,200,66,0.4) !important;
}
[data-testid="stSidebar"] .stButton > button:focus,
[data-testid="stSidebar"] .stButton > button:active {
    box-shadow: none !important; outline: none !important;
    background: rgba(255,255,255,0.05) !important;
}

/* Club row */
.club-row .stButton > button {
    font-size: 11px !important; font-weight: 700 !important;
    color: rgba(255,255,255,0.75) !important;
    text-transform: uppercase !important; letter-spacing: 0.5px !important;
    padding: 10px 20px !important;
}

/* Player rows */
.player-row .stButton > button {
    color: rgba(255,255,255,0.4) !important;
    padding: 6px 20px 6px 34px !important;
    font-size: 13px !important; font-style: italic !important;
}
.player-active .stButton > button {
    color: #ffffff !important; font-style: normal !important;
    font-weight: 600 !important;
    border-left: 2px solid #f5c842 !important;
    background: rgba(245,200,66,0.06) !important;
}

/* Add player row */
.add-player-row .stButton > button {
    color: rgba(245,200,66,0.4) !important;
    font-size: 11px !important;
    padding: 5px 20px 8px 34px !important;
    border-top: 1px solid rgba(255,255,255,0.04) !important;
}
.add-player-row .stButton > button:hover {
    color: #f5c842 !important;
    background: rgba(245,200,66,0.05) !important;
}

/* New academy button */
.new-acad-row .stButton > button {
    color: rgba(255,255,255,0.35) !important;
    font-size: 11px !important; padding: 8px 20px !important;
    border: 1px dashed rgba(255,255,255,0.12) !important;
    border-radius: 8px !important;
    margin: 4px 16px !important;
    width: calc(100% - 32px) !important;
    text-align: center !important;
}
.new-acad-row .stButton > button:hover {
    color: rgba(255,255,255,0.8) !important;
    border-color: rgba(255,255,255,0.3) !important;
    background: rgba(255,255,255,0.04) !important;
}

/* Mode buttons */
.mode-btn-wrap { padding: 8px 16px 4px; display: flex; gap: 6px; }
[data-testid="stSidebar"] .stButton > button:focus { outline: none !important; box-shadow: none !important; }

/* Sign out */
.signout-wrap { padding: 8px 16px 16px; }
.signout-wrap .stButton > button {
    background: rgba(255,255,255,0.04) !important;
    color: rgba(255,255,255,0.3) !important;
    border: 1px solid rgba(255,255,255,0.08) !important;
    border-radius: 8px !important;
    font-size: 11px !important; letter-spacing: 1px !important;
    text-transform: uppercase !important;
    padding: 8px 16px !important;
    margin: 0 !important; width: 100% !important;
    text-align: center !important;
}
.signout-wrap .stButton > button:hover {
    background: rgba(239,68,68,0.1) !important;
    color: #fca5a5 !important;
    border-color: rgba(239,68,68,0.2) !important;
}

.sb-meta {
    padding: 6px 20px 12px;
    font-size: 10px; color: rgba(255,255,255,0.18);
    line-height: 1.8;
}

/* ── MAIN CONTENT ── */

/* Page header */
.page-header {
    padding: 32px 0 24px;
    border-bottom: 1px solid #e5e7ef;
    margin-bottom: 28px;
    display: flex; align-items: flex-end;
    justify-content: space-between;
}
.page-title {
    font-family: 'Playfair Display', serif !important;
    font-size: 42px; font-weight: 700;
    color: #04080f; line-height: 1; margin: 0;
    letter-spacing: -1px;
}
.page-meta {
    font-size: 10px; font-weight: 600;
    color: #6b7280; letter-spacing: 3px;
    text-transform: uppercase; margin-top: 8px;
}
.page-badge {
    display: inline-flex; align-items: center; gap: 6px;
    background: #f0fdf4; border: 1px solid #bbf7d0;
    border-radius: 20px; padding: 4px 12px;
    font-size: 11px; font-weight: 600; color: #16a34a;
}

/* Metric cards */
.metric-grid { display: grid; grid-template-columns: repeat(6,1fr); gap: 12px; margin-bottom: 8px; }
.metric-card {
    background: #ffffff;
    border: 1px solid #e5e7ef;
    border-radius: 12px; padding: 16px;
    box-shadow: 0 1px 3px rgba(0,0,0,0.04), 0 4px 12px rgba(0,0,0,0.02);
    transition: box-shadow 0.2s;
}
.metric-card:hover { box-shadow: 0 4px 16px rgba(0,0,0,0.08); }
.metric-label {
    font-size: 9px; font-weight: 700;
    color: #9ca3af; letter-spacing: 2px;
    text-transform: uppercase; margin-bottom: 8px;
}
.metric-value {
    font-size: 22px; font-weight: 700; color: #04080f; line-height: 1;
}
.metric-card.highlight { border-top: 3px solid #f5c842; }
.metric-card.positive { border-top: 3px solid #16a34a; }
.metric-card.neutral { border-top: 3px solid #3b82f6; }

/* Section titles */
.section-title {
    font-family: 'Playfair Display', serif !important;
    font-size: 18px; font-weight: 600; color: #04080f;
    margin: 32px 0 14px; padding-bottom: 12px;
    border-bottom: 1px solid #e5e7ef;
    display: flex; align-items: center; justify-content: space-between;
}

/* Report */
.report-wrapper {
    background: #ffffff; border: 1px solid #e5e7ef;
    border-radius: 16px; padding: 44px 52px;
    box-shadow: 0 2px 8px rgba(0,0,0,0.04);
}
.report-header {
    border-bottom: 2px solid #04080f; padding-bottom: 20px; margin-bottom: 28px;
}
.report-title {
    font-family: 'Playfair Display', serif !important;
    font-size: 11px; font-weight: 600; color: #f5c842;
    letter-spacing: 4px; text-transform: uppercase; margin-bottom: 8px;
}
.report-player-name {
    font-family: 'Playfair Display', serif !important;
    font-size: 28px; font-weight: 700; color: #04080f;
}
.report-section-label {
    font-size: 9px; font-weight: 700; color: #3b82f6;
    letter-spacing: 3px; text-transform: uppercase;
    margin: 28px 0 10px; padding-bottom: 8px;
    border-bottom: 1px solid #f3f4f6;
}
.report-para {
    font-size: 14px; color: #374151; line-height: 1.85;
    margin-bottom: 4px;
}
.exec-summary {
    background: #fffbeb; border-left: 4px solid #f5c842;
    border-radius: 0 12px 12px 0; padding: 20px 24px;
    margin-bottom: 28px;
}
.exec-summary .report-para { color: #1c1917; font-weight: 500; }

/* Coach notes */
.coach-notes-wrap {
    background: #fffbeb; border: 1px solid #fde68a;
    border-radius: 12px; padding: 20px 24px; margin-top: 20px;
}
.coach-notes-label {
    font-size: 9px; font-weight: 700; color: #92400e;
    letter-spacing: 3px; text-transform: uppercase; margin-bottom: 10px;
}

/* Action buttons */
.btn-row { display: flex; gap: 8px; margin-bottom: 20px; }
.export-btn .stButton > button {
    background: #04080f !important; color: #ffffff !important;
    border: none !important; border-radius: 8px !important;
    font-size: 10px !important; font-weight: 700 !important;
    letter-spacing: 1px !important; text-transform: uppercase !important;
    padding: 8px 16px !important; box-shadow: none !important;
    width: auto !important;
}
.export-btn .stButton > button:hover { background: #1a3a8a !important; }
.edit-btn .stButton > button {
    background: #7c3aed !important; color: #fff !important;
    border: none !important; border-radius: 8px !important;
    font-size: 10px !important; font-weight: 700 !important;
    letter-spacing: 1px !important; text-transform: uppercase !important;
    padding: 8px 16px !important; box-shadow: none !important; width: auto !important;
}
.save-btn .stButton > button {
    background: #16a34a !important; color: #fff !important;
    border: none !important; border-radius: 8px !important;
    font-size: 10px !important; font-weight: 700 !important;
    letter-spacing: 1px !important; text-transform: uppercase !important;
    padding: 8px 16px !important; box-shadow: none !important; width: auto !important;
}
.gen-btn .stButton > button {
    background: linear-gradient(135deg, #16a34a, #15803d) !important;
    color: #ffffff !important; border: none !important;
    border-radius: 10px !important; font-weight: 700 !important;
    font-size: 13px !important; letter-spacing: 0.5px !important;
    padding: 14px 28px !important; box-shadow: 0 4px 16px rgba(22,163,74,0.3) !important;
    width: auto !important;
}
.regen-btn .stButton > button {
    background: rgba(245,200,66,0.1) !important; color: #92400e !important;
    border: 1px solid rgba(245,200,66,0.4) !important;
    border-radius: 8px !important; font-size: 10px !important;
    font-weight: 700 !important; letter-spacing: 1px !important;
    text-transform: uppercase !important; padding: 8px 16px !important;
    box-shadow: none !important; width: auto !important;
}
.stDownloadButton > button {
    background: #dc2626 !important; color: #fff !important;
    font-size: 10px !important; letter-spacing: 1px !important;
    text-transform: uppercase !important; padding: 8px 16px !important;
    border-radius: 8px !important; border: none !important;
    font-weight: 700 !important; white-space: nowrap !important;
    width: auto !important; min-width: 0 !important; max-width: 160px !important;
}

/* Tables */
.stDataFrame { border: 1px solid #e5e7ef !important; border-radius: 12px !important; }

/* Text areas */
.stTextArea textarea {
    font-family: 'Outfit', sans-serif !important;
    font-size: 13px !important; color: #374151 !important;
    border: 1px solid #e5e7ef !important; border-radius: 10px !important;
    line-height: 1.8 !important;
}

/* Welcome cards */
.welcome-grid { display: grid; grid-template-columns: repeat(3,1fr); gap: 16px; margin-top: 24px; }
.welcome-card {
    background: #ffffff; border: 1px solid #e5e7ef;
    border-radius: 16px; padding: 28px 24px;
    text-align: center;
    box-shadow: 0 1px 3px rgba(0,0,0,0.04);
    transition: all 0.2s;
}
.welcome-card:hover {
    box-shadow: 0 8px 24px rgba(0,0,0,0.08);
    transform: translateY(-2px);
}
.welcome-icon { font-size: 32px; margin-bottom: 12px; }
.welcome-title { font-size: 15px; font-weight: 700; color: #04080f; margin-bottom: 6px; }
.welcome-desc { font-size: 12px; color: #9ca3af; line-height: 1.6; }

/* Upload banner */
.upload-banner {
    background: #eff6ff; border: 1px solid #bfdbfe;
    border-radius: 10px; padding: 16px 20px; margin-bottom: 20px;
}

/* No select */
.no-select { padding: 80px 24px; text-align: center; color: #9ca3af; }

/* Pro badge */
.badge-pro {
    display: inline-block; background: rgba(245,200,66,0.1);
    border: 1px solid rgba(245,200,66,0.3); color: #92400e;
    font-size: 8px; font-weight: 700; letter-spacing: 2px;
    padding: 2px 7px; border-radius: 4px; margin-left: 6px;
    text-transform: uppercase;
}
.badge-beta {
    display: inline-block; background: rgba(59,130,246,0.1);
    border: 1px solid rgba(59,130,246,0.2); color: #1d4ed8;
    font-size: 8px; font-weight: 700; letter-spacing: 2px;
    padding: 2px 7px; border-radius: 4px; margin-left: 6px;
    text-transform: uppercase;
}

/* Scrollbar */
::-webkit-scrollbar { width: 4px; height: 4px; }
::-webkit-scrollbar-track { background: transparent; }
::-webkit-scrollbar-thumb { background: #d1d5db; border-radius: 2px; }

div[role="radiogroup"] { display: none; }
.sb-dots .stButton > button {
    background: transparent !important; border: none !important;
    color: rgba(255,255,255,0.25) !important; font-size: 14px !important;
    padding: 2px 5px !important; width: auto !important;
    min-width: 0 !important; box-shadow: none !important; margin: 0 !important;
}
.sb-dots .stButton > button:hover { color: rgba(255,255,255,0.7) !important; }
.dot-edit .stButton > button {
    color: rgba(245,200,66,0.8) !important; font-size: 11px !important;
    padding: 4px 10px !important; width: 100% !important; text-align: left !important;
    border-radius: 5px !important; background: rgba(245,200,66,0.06) !important;
    border: 1px solid rgba(245,200,66,0.15) !important; margin-bottom: 2px !important;
}
.dot-del .stButton > button {
    color: rgba(239,68,68,0.8) !important; font-size: 11px !important;
    padding: 4px 10px !important; width: 100% !important; text-align: left !important;
    border-radius: 5px !important; background: rgba(239,68,68,0.06) !important;
    border: 1px solid rgba(239,68,68,0.15) !important;
}

/* ── FILE UPLOADER - minimal clean ── */
[data-testid="stFileUploader"] {
    width: 100% !important;
}
[data-testid="stFileUploader"] section {
    background: #f8f9fc !important;
    border: 1.5px dashed #d1d5db !important;
    border-radius: 10px !important;
    padding: 12px 16px !important;
    min-height: 0 !important;
}
[data-testid="stFileUploader"] section:hover {
    border-color: #1a3a8a !important;
    background: #f0f4ff !important;
}
/* Hide the verbose drag-and-drop text */
[data-testid="stFileUploaderDropzoneInstructions"] {
    display: none !important;
}
/* The browse button */
[data-testid="stFileUploaderDropzone"] button,
[data-testid="baseButton-secondary"] {
    background: #04080f !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 100px !important;
    font-family: 'Outfit', sans-serif !important;
    font-size: 11px !important;
    font-weight: 600 !important;
    padding: 7px 20px !important;
    cursor: pointer !important;
    white-space: nowrap !important;
    width: auto !important;
    display: inline-block !important;
}
[data-testid="stFileUploaderDropzone"] button:hover {
    background: #1a3a8a !important;
}
/* Uploaded file name display */
[data-testid="stFileUploader"] [data-testid="stMarkdownContainer"] p {
    font-size: 11px !important;
    color: #6b7280 !important;
}
</style>
""", unsafe_allow_html=True)
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_CENTER
    import anthropic
except ImportError as e:
    st.error(f"Missing library: {e}")
    st.stop()

def get_api_key():
    try:
        return st.secrets["ANTHROPIC_KEY"]
    except Exception:
        return os.getenv("ANTHROPIC_KEY")

def ai_report(prompt):
    claude = anthropic.Anthropic(api_key=get_api_key())
    r = claude.messages.create(
        model="claude-opus-4-5", max_tokens=5000,
        messages=[{"role": "user", "content": prompt}]
    )
    return r.content[0].text

def ai_clean_data(raw_text, data_type="youth"):
    """Parse any unstructured data into platform format using Claude."""
    api_key = get_api_key()
    if not api_key:
        return "[]"
    client = anthropic.Anthropic(api_key=api_key)
    if data_type == "youth":
        prompt = f"""You are a football data parser. Convert the raw data below into a JSON array for a youth player tracking platform.

The data can be in ANY format: match notes, WhatsApp messages, copied tables, handwritten notes, CSV text, anything.
Extract every player you can find. Use sensible defaults for missing fields.

Return ONLY a valid JSON array. Each object must have:
name, date_of_birth (YYYY-MM-DD or null), position (Striker/Right Winger/Left Winger/Attacking Midfielder/Central Midfielder/Defensive Midfielder/Right Back/Left Back/Centre Back/Goalkeeper),
club, dominant_foot (Right/Left), age_group (U13/U14/U15/U16/U17/U18), nationality,
session_date (YYYY-MM-DD or null), session_type (match/training), minutes_played (integer),
distance_covered_km (float), sprint_count (integer), top_speed_kmh (float),
passes_completed (integer), passes_attempted (integer), dribbles_completed (integer),
defensive_actions (integer), goals (integer), assists (integer), chances_created (integer),
tackles_won (integer), coachability_rating (1-10), attitude_score (1-10),
consistency_rating (1-10), coach_notes (string)

Defaults: 7 for ratings, 90 for minutes, 8.0 for distance, right foot, U16.
No markdown. No explanation. Return ONLY the JSON array.

Raw data:
{raw_text}"""
    else:
        prompt = f"""You are a football data parser for a professional scouting platform.
Convert the raw data below into a JSON array of professional player records.

The data can be in ANY format: copied from websites, spreadsheets, PDF reports, stats tables, rough notes.
If the data is about one player (e.g. "Bruno Fernandes Stats"), create one record for that player.

IMPORTANT RULES:
- name = the PLAYER's name (e.g. "Bruno Fernandes"), NOT the club name
- team = the CLUB the player plays for (e.g. "Manchester United")
- goals = actual goals scored (the "Goals" field) — NOT xG
- assists = actual assists (the "Assists" field) — NOT xA
- xg = expected goals (the "XG" or "xG" field) — separate from goals
- xa = expected assists (the "XA" or "xA" field) — separate from assists
- minutes = total minutes played (the "Minutes Played" or "Min" field)
- appearances = number of matches played (the "Appearances" or "Apps" field)
- Passes completed is NOT assists — it's passing volume

Return ONLY a valid JSON array. Each object must have:
name (string), team (string), position (Forward/Midfielder/Defender/Goalkeeper),
nationality (string), age (integer), season (string),
appearances (integer), minutes (integer), goals (integer), assists (integer),
goals_per_90 (float), assists_per_90 (float),
yellow_cards (integer), red_cards (integer), clean_sheets (integer),
xg (float), xa (float)

Use 0 for missing numeric values. No markdown. No explanation. Return ONLY the JSON array.

Raw data:
{raw_text}"""
    r = client.messages.create(
        model="claude-opus-4-5",
        max_tokens=4000,
        messages=[{{"role": "user", "content": prompt}}]
    )
    return r.content[0].text


def to_pdf(name, meta, text, is_pro=False):
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=22*mm, rightMargin=22*mm, topMargin=20*mm, bottomMargin=20*mm)
    navy=HexColor('#04080f'); blue=HexColor('#1a3a8a'); gold=HexColor('#f5c842')
    bc=HexColor('#374151'); lc=HexColor('#9ca3af'); dc=HexColor('#e5e7ef')
    s = {
        'logo': ParagraphStyle('l', fontName='Helvetica-Bold', fontSize=20, textColor=navy, spaceAfter=2, leading=24),
        'tag':  ParagraphStyle('t', fontName='Helvetica', fontSize=7, textColor=lc, spaceAfter=8, leading=10),
        'name': ParagraphStyle('n', fontName='Helvetica-Bold', fontSize=16, textColor=navy, spaceAfter=3, leading=20),
        'meta': ParagraphStyle('m', fontName='Helvetica', fontSize=8, textColor=lc, spaceAfter=10, leading=12),
        'head': ParagraphStyle('h', fontName='Helvetica-Bold', fontSize=8, textColor=blue, spaceAfter=4, spaceBefore=14, leading=12),
        'body': ParagraphStyle('b', fontName='Helvetica', fontSize=9, textColor=bc, spaceAfter=3, leading=15),
        'foot': ParagraphStyle('f', fontName='Helvetica', fontSize=7, textColor=lc, alignment=TA_CENTER),
    }
    tag = "PRO ANALYSIS — ENGLISH PREMIER LEAGUE" if is_pro else "YOUTH TALENT INTELLIGENCE"
    story = [
        Paragraph("SCOUT IQ·FC", s['logo']), Spacer(1,2*mm),
        Paragraph(tag, s['tag']),
        HRFlowable(width="100%", thickness=1.5, color=gold, spaceAfter=6*mm),
        Paragraph(name.upper(), s['name']), Paragraph(meta, s['meta']),
        HRFlowable(width="100%", thickness=0.3, color=dc, spaceAfter=8*mm),
    ]
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            story.append(Spacer(1,3*mm))
        elif line[0].isdigit() and '.' in line[:3]:
            story.append(HRFlowable(width="100%", thickness=0.2, color=dc, spaceBefore=4*mm, spaceAfter=3*mm))
            story.append(Paragraph(line.upper(), s['head']))
        else:
            clean = line.replace('**','').replace('--','').replace('- ','').replace('#','')
            if clean:
                story.append(Paragraph(clean, s['body']))
    story += [Spacer(1,10*mm), HRFlowable(width="100%",thickness=0.5,color=gold),
              Spacer(1,3*mm), Paragraph(f"Scout IQ·FC  ·  {name}  ·  Confidential", s['foot'])]
    doc.build(story)
    buf.seek(0)
    return buf

def to_word(name, meta, text, is_pro=False):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)
    t = doc.add_heading('SCOUT IQ·FC', 0)
    t.runs[0].font.color.rgb = RGBColor(4,8,15)
    t.runs[0].font.size = Pt(20); t.runs[0].font.name = 'Calibri'
    sub = doc.add_paragraph("YOUTH TALENT INTELLIGENCE" if not is_pro else "PRO ANALYSIS")
    sub.runs[0].font.size = Pt(8); sub.runs[0].font.color.rgb = RGBColor(156,163,175)
    doc.add_paragraph()
    n = doc.add_heading(name.upper(), 1)
    n.runs[0].font.color.rgb = RGBColor(4,8,15); n.runs[0].font.size = Pt(16); n.runs[0].font.name = 'Calibri'
    m = doc.add_paragraph(meta)
    m.runs[0].font.size = Pt(9); m.runs[0].font.color.rgb = RGBColor(156,163,175)
    doc.add_paragraph()
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line[0].isdigit() and '.' in line[:3]:
            h = doc.add_heading(line, 2)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(26,58,138)
                h.runs[0].font.size = Pt(10); h.runs[0].font.name = 'Calibri'; h.runs[0].bold = True
        else:
            clean = line.replace('**','').replace('--','').replace('- ','').replace('#','')
            if clean:
                p = doc.add_paragraph(clean)
                if p.runs:
                    p.runs[0].font.size = Pt(10); p.runs[0].font.name = 'Calibri'
                    p.runs[0].font.color.rgb = RGBColor(55,65,81)
    fp = doc.sections[0].footer.paragraphs[0]
    fp.text = f"Scout IQ·FC   ·   {name}   ·   Confidential"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

def render_report(text, pid, pname, meta, is_pro=False):
    cursor.execute("SELECT edited_report, coach_notes FROM report_edits WHERE player_id=? ORDER BY edited_at DESC LIMIT 1", (pid,))
    row = cursor.fetchone()
    active = row[0] if row else text
    notes = row[1] if row and row[1] else ""
    ekey = f"edit_{pid}_{is_pro}"
    if ekey not in st.session_state:
        st.session_state[ekey] = False
    # Action bar — small buttons
    c1,c2,c3,c4,c5 = st.columns([1,1,1,1,6])
    with c1:
        st.markdown('<div class="edit-btn">', unsafe_allow_html=True)
        if st.button("Edit", key=f"eb_{pid}_{is_pro}"):
            st.session_state[ekey] = not st.session_state[ekey]; st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
    with c2:
        st.markdown('<div class="export-btn">', unsafe_allow_html=True)
        if st.button("PDF", key=f"pb_{pid}_{is_pro}"):
            st.session_state[f"sp_{pid}_{is_pro}"] = True
        st.markdown('</div>', unsafe_allow_html=True)
    with c3:
        st.markdown('<div class="export-btn">', unsafe_allow_html=True)
        if st.button("Word", key=f"wb_{pid}_{is_pro}"):
            st.session_state[f"sw_{pid}_{is_pro}"] = True
        st.markdown('</div>', unsafe_allow_html=True)
    with c4:
        st.markdown('<div class="regen-btn">', unsafe_allow_html=True)
        if st.button("↺ Redo", key=f"rg_{pid}_{is_pro}"):
            st.session_state[f"confirm_regen_{pid}_{is_pro}"] = True
        st.markdown('</div>', unsafe_allow_html=True)
    if st.session_state.get(f"confirm_regen_{pid}_{is_pro}"):
        st.warning("This will replace the existing report. Confirm?")
        cc1, _ = st.columns([1,5])
        with cc1:
            if st.button("Yes, regenerate", key=f"cy_{pid}_{is_pro}"):
                cursor.execute("DELETE FROM epl_reports WHERE player_id=?" if is_pro else "DELETE FROM reports WHERE player_id=?", (pid,))
                conn.commit()
                st.session_state[f"confirm_regen_{pid}_{is_pro}"] = False
                st.rerun()
    if st.session_state.get(f"sp_{pid}_{is_pro}"):
        buf = to_pdf(pname, meta, active, is_pro)
        st.download_button("⬇ Download PDF", buf, f"ScoutIQFC_{pname.replace(' ','_')}.pdf", "application/pdf", key=f"dp_{pid}_{is_pro}")
    if st.session_state.get(f"sw_{pid}_{is_pro}"):
        buf = to_word(pname, meta, active, is_pro)
        st.download_button("⬇ Download Word", buf, f"ScoutIQFC_{pname.replace(' ','_')}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"dw_{pid}_{is_pro}")
    if st.session_state[ekey]:
        st.markdown('<div class="section-title">Edit Report</div>', unsafe_allow_html=True)
        edited = st.text_area("Report", value=active, height=500, key=f"ea_{pid}_{is_pro}", label_visibility="collapsed")
        st.markdown('<div class="coach-notes-wrap"><div class="coach-notes-label">Coach Annotations</div>', unsafe_allow_html=True)
        new_notes = st.text_area("Notes", value=notes, height=120, key=f"na_{pid}_{is_pro}", label_visibility="collapsed")
        st.markdown('</div>', unsafe_allow_html=True)
        st.markdown('<div class="save-btn">', unsafe_allow_html=True)
        if st.button("Save Changes", key=f"sv_{pid}_{is_pro}"):
            cursor.execute("INSERT INTO report_edits (player_id,edited_report,coach_notes) VALUES (?,?,?)", (pid, edited, new_notes))
            conn.commit(); st.session_state[ekey] = False; st.success("Saved"); st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
    else:
        # ── REPORT DISPLAY — dark navy, gold accents, no white boxes ──
        report_label = "Pro Performance Audit" if is_pro else "Youth Scouting Report"
        html = f'<div style="background:#04080f;border:1px solid rgba(245,200,66,0.2);border-radius:16px;padding:40px 48px;margin-top:8px;box-shadow:0 4px 32px rgba(0,0,0,0.3);"><div style="font-size:10px;font-weight:700;color:#f5c842;letter-spacing:4px;text-transform:uppercase;margin-bottom:12px;">Scout IQ·FC — {report_label}</div><div style="font-family:Georgia,serif;font-size:36px;font-weight:700;color:#ffffff;line-height:1.1;margin-bottom:6px;">{pname}</div><div style="font-size:12px;color:rgba(255,255,255,0.35);margin-bottom:0;">{meta}</div><div style="border-top:1px solid rgba(245,200,66,0.25);margin:20px 0;"></div>'
        in_exec = False
        for line in active.split("\n"):
            line = line.strip()
            if not line:
                if in_exec:
                    html += '</div>'
                    in_exec = False
                html += '<div style="height:6px"></div>'
            elif line.upper().startswith("EXECUTIVE SUMMARY"):
                html += '<div style="background:rgba(245,200,66,0.06);border-left:3px solid #f5c842;border-radius:0 10px 10px 0;padding:16px 20px;margin-bottom:24px;"><div style="font-size:9px;font-weight:700;color:#f5c842;letter-spacing:3px;text-transform:uppercase;margin-bottom:10px;">Executive Summary</div>'
                in_exec = True
            elif line[0].isdigit() and "." in line[:3]:
                if in_exec:
                    html += '</div>'
                    in_exec = False
                html += f'<div style="font-size:9px;font-weight:700;color:#f5c842;letter-spacing:3px;text-transform:uppercase;margin:28px 0 10px;padding-bottom:8px;border-bottom:1px solid rgba(255,255,255,0.07);">{line}</div>'
            else:
                clean = line.replace("**","").replace("--","").replace("#","")
                if clean:
                    color = "rgba(255,255,255,0.75)" if in_exec else "rgba(255,255,255,0.6)"
                    weight = "500" if in_exec else "400"
                    html += f'<div style="font-size:14px;color:{color};line-height:1.9;margin-bottom:2px;font-weight:{weight};">{clean}</div>'
        if in_exec:
            html += '</div>'
        html += '</div>'
        st.markdown(html, unsafe_allow_html=True)
        if notes:
            st.markdown(
                f'<div class="coach-notes-wrap" style="margin-top:12px;">                <div class="coach-notes-label">Coach Annotations</div>                <div style="font-size:13px;color:#92400e;line-height:1.8;">{notes}</div></div>',
                unsafe_allow_html=True
            )

def sg(row, idx, default=0):
    try:
        v = row[idx]; return v if v is not None else default
    except IndexError:
        return default

def build_youth_prompt(player, sessions):
    sc = len(sessions)
    avg_dist = round(sum(s[5] for s in sessions)/sc, 2) if sc else 0
    avg_spr = round(sum(s[6] for s in sessions)/sc, 1) if sc else 0
    peak_spd = max(s[7] for s in sessions) if sc else 0
    total_pc = sum(s[8] for s in sessions); total_pa = sum(s[9] for s in sessions)
    pass_pct = round(total_pc/total_pa*100, 1) if total_pa > 0 else 0
    total_goals = sum(s[12] for s in sessions); total_assists = sum(s[13] for s in sessions)
    total_def = sum(s[11] for s in sessions)
    avg_coach = round(sum(s[16] for s in sessions)/sc, 1) if sc else 0
    avg_att = round(sum(s[17] for s in sessions)/sc, 1) if sc else 0
    avg_cons = round(sum(s[18] for s in sessions)/sc, 1) if sc else 0
    total_mins = sum(s[4] for s in sessions)
    goals_p90 = round(total_goals/(total_mins/90), 2) if total_mins > 0 else 0
    ast_p90 = round(total_assists/(total_mins/90), 2) if total_mins > 0 else 0
    def_p90 = round(total_def/(total_mins/90), 2) if total_mins > 0 else 0
    sessions_text = "\n".join([
        f"- {s[2]} ({s[3]}): {s[4]}mins, {s[5]}km, {s[6]} sprints, {s[7]}km/h top speed, "
        f"passes {s[8]}/{s[9]} ({round(s[8]/s[9]*100) if s[9]>0 else 0}%), drib {s[10]}, "
        f"def {s[11]}, G{s[12]} A{s[13]}, tackles {s[15]}. "
        f"Coach {s[16]}/10 Att {s[17]}/10 Cons {s[18]}/10. Notes: {s[19]}"
        for s in sessions
    ])
    return f"""You are a senior scout and player development analyst writing for an academy director and technical staff. Every claim must cite a specific number. No generic phrases. No bullet points. Full sentences only. Maximum 4 printed pages.

PLAYER: {player[1]} | {player[4]} | {player[3]} | {player[6]} | DOB: {player[2]} | {player[7] or 'Unknown'} | {player[5]} foot
SESSIONS: {sc} | Pass%: {pass_pct}% | G/90: {goals_p90} | A/90: {ast_p90} | Def/90: {def_p90}
Avg dist: {avg_dist}km | Avg sprints: {avg_spr} | Peak speed: {peak_spd}km/h
Coachability: {avg_coach}/10 | Attitude: {avg_att}/10 | Consistency: {avg_cons}/10
SESSION LOG: {sessions_text}

Write the complete report:

EXECUTIVE SUMMARY
4 sentences: strongest metric with number, biggest weakness with number, trajectory trend, one actionable recommendation.

1. PERFORMANCE RATING
Score/10 with two data points. Above/at/below expectation for age and position.

2. TECHNICAL PROFILE
Pass %, dribble output, defensive actions/90, goal contributions/90. Two sentences on quality vs age/position expectations. Position-appropriate verdict.

3. PHYSICAL PROFILE
Exact avg distance, sprints, peak speed. Elite/adequate/below for {player[3]} {player[4]}. Flag lowest output session with date and likely cause.

4. MENTAL AND ATTITUDE PROFILE
Exact averages for all three scores. Trend direction. What it means for development. One specific coaching instruction.

5. BEST POSITION NOW AND FUTURE
Current optimal position from data. One alternative with two sentence justification. State if being played out of position.

6. SHORT TERM TRAINING PRIORITIES (4-8 weeks)
Three interventions: current metric, target, specific session format.

7. SEASON DEVELOPMENT TARGETS
Three measurable targets with baselines and timeframes. One unconventional recommendation.

8. CEILING AND LONG TERM PATHWAY
Level: local/regional/national/professional. Three data point justification. One paragraph each for trajectory with and without intervention.

9. LOAD AND INJURY WATCH
Short/medium/long term assessment. If no concerns, state with data.

10. COACHING AND MANAGEMENT GUIDE
Short term feedback approach. Season minute management. Long term psychological readiness verdict.

11. SESSION DESIGN IDEAS
Two sessions: name, setup (2 sentences), metric targeted, why it suits this player. Keep concise.

12. SCOUT VERDICT
Category: Continue Monitoring / Increase Development Investment / Priority Development Case / Recommend for Promotion. Two sentences with specific numbers. One memorable verbal description.

Complete all 12 sections. No truncation. No filler."""

# ── LOAD DATA ──
cursor.execute("SELECT id,name,position,club,age_group,nationality,dominant_foot FROM players")
all_players = cursor.fetchall()
# Beta accounts only see one demo academy
account_type = st.session_state.get("account_type", "admin")
BETA_ACADEMY = "Riverside FC Academy"
if account_type == "beta":
    players = [p for p in all_players if (p[3] or "Unassigned") == BETA_ACADEMY]
else:
    players = all_players
clubs = {}
for p in players:
    club = p[3] or "Unassigned"
    if club not in clubs: clubs[club] = []
    clubs[club].append(p)
cursor.execute("SELECT DISTINCT team FROM epl_players ORDER BY team")
epl_teams = [r[0] for r in cursor.fetchall()]
for k,v in {
    "selected_player_id": None,
    "expanded_clubs": [], "expanded_epl_teams": [],
    "selected_epl_player_id": None, "mode": "youth",
    "show_add_player": False, "show_add_academy": False, "selected_academy": None,
    "add_player_club": None, "show_ai_import_academy": None, "ai_parsed_players": None
}.items():
    if k not in st.session_state: st.session_state[k] = v

# ══════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════
with st.sidebar:
    st.markdown("""
    <div class="sb-logo">
        <div class="sb-logo-name">Scout IQ·<span class="fc">FC</span></div>
        <div class="sb-logo-tag">Where Talent Is Unearthed</div>
    </div>
    """, unsafe_allow_html=True)
    if st.session_state.get("account_type") == "beta":
        st.markdown('<div style="background:rgba(245,200,66,0.1);border:1px solid rgba(245,200,66,0.25);border-radius:6px;padding:6px 12px;margin:6px 16px;font-size:10px;color:rgba(245,200,66,0.8);text-align:center;font-weight:600;letter-spacing:1px;">BETA ACCESS</div>', unsafe_allow_html=True)
    st.markdown('<div class="sb-section">Mode</div>', unsafe_allow_html=True)
    m1,m2 = st.columns(2)
    with m1:
        if st.button("Youth", key="ny", use_container_width=True):
            st.session_state.mode="youth"; st.rerun()
    with m2:
        if st.button("Pro Data", key="np", use_container_width=True):
            st.session_state.mode="pro"; st.rerun()
    st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)
    if st.session_state.mode == "youth":
        srch = st.text_input("s", placeholder="Search players...", label_visibility="collapsed")
        st.markdown('<div class="sb-section">Academies <span class="badge-beta">Beta</span></div>', unsafe_allow_html=True)
        st.markdown('<div class="new-acad-row">', unsafe_allow_html=True)
        if st.button("＋ New Academy", key="new_acad", use_container_width=True):
            st.session_state.show_add_academy = True
            st.session_state.show_add_player = False
            st.session_state.selected_player_id = None
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
        for club, cplayers in clubs.items():
            filt = [p for p in cplayers if srch.lower() in p[1].lower()] if srch else cplayers
            if not filt: continue
            exp = club in st.session_state.expanded_clubs or bool(srch)
            ac1, ac2 = st.columns([8, 1])
            with ac1:
                st.markdown('<div class="club-row">', unsafe_allow_html=True)
                if st.button(f"{'▾' if exp else '▸'}  {club[:22]}", key=f"cl_{club}", use_container_width=True):
                    if exp and not srch: st.session_state.expanded_clubs.remove(club)
                    elif not exp: st.session_state.expanded_clubs.append(club)
                    st.rerun()
                st.markdown('</div>', unsafe_allow_html=True)
            with ac2:
                st.markdown('<div class="sb-dots">', unsafe_allow_html=True)
                if st.button("⋮", key=f"adots_{club}"):
                    cur = st.session_state.get("open_menu")
                    st.session_state["open_menu"] = None if cur == f"acad_{club}" else f"acad_{club}"
                    st.rerun()
                st.markdown('</div>', unsafe_allow_html=True)
            if st.session_state.get("open_menu") == f"acad_{club}":
                st.markdown('<div style="background:rgba(255,255,255,0.04);border:1px solid rgba(255,255,255,0.09);border-radius:8px;padding:6px 8px;margin:2px 14px 4px;">', unsafe_allow_html=True)
                st.markdown('<div class="dot-edit">', unsafe_allow_html=True)
                if st.button("✏  Edit Academy", key=f"edit_acad_{club}"):
                    st.session_state["open_menu"] = None
                    st.session_state.show_edit_academy = True
                    st.session_state.edit_academy_name = club
                    st.session_state.selected_player_id = None
                    st.rerun()
                st.markdown('</div><div class="dot-del">', unsafe_allow_html=True)
                if st.button("✕  Delete Academy", key=f"del_acad_{club}"):
                    st.session_state["open_menu"] = None
                    st.session_state["confirm_del_acad"] = club
                    st.rerun()
                st.markdown('</div></div>', unsafe_allow_html=True)
            if st.session_state.get("confirm_del_acad") == club:
                st.markdown(f'<div style="background:rgba(239,68,68,0.08);border:1px solid rgba(239,68,68,0.2);border-radius:6px;padding:7px 12px;margin:2px 14px;font-size:11px;color:#fca5a5;">Delete {club}?</div>', unsafe_allow_html=True)
                cc1, cc2 = st.columns(2)
                with cc1:
                    if st.button("Yes", key=f"cda_{club}"):
                        for dp in cplayers:
                            cursor.execute("DELETE FROM sessions WHERE player_id=?", (dp[0],))
                            cursor.execute("DELETE FROM reports WHERE player_id=?", (dp[0],))
                        cursor.execute("DELETE FROM players WHERE club=?", (club,))
                        conn.commit()
                        st.session_state["confirm_del_acad"] = None
                        st.session_state.selected_player_id = None
                        st.rerun()
                with cc2:
                    if st.button("No", key=f"nda_{club}"):
                        st.session_state["confirm_del_acad"] = None; st.rerun()
            if exp:
                st.markdown('<div class="player-row">', unsafe_allow_html=True)
                if st.button(f"   📊  Academy Overview", key=f"acad_view_{club}", use_container_width=True):
                    st.session_state.selected_academy = club
                    st.session_state.selected_player_id = None
                    st.session_state.show_add_player = False
                    st.rerun()
                st.markdown('</div>', unsafe_allow_html=True)
                for p in filt:
                    active = st.session_state.selected_player_id == p[0]
                    pc1, pc2 = st.columns([8, 1])
                    with pc1:
                        st.markdown(f'<div class="{"player-active" if active else "player-row"}">', unsafe_allow_html=True)
                        if st.button(f"{'▸  ' if active else '   '}{p[1][:20]}", key=f"yp_{p[0]}", use_container_width=True):
                            st.session_state.selected_player_id = p[0]
                            st.session_state.show_add_player = False
                            st.session_state.show_add_academy = False
                            st.session_state["open_menu"] = None
                            st.rerun()
                        st.markdown('</div>', unsafe_allow_html=True)
                    with pc2:
                        st.markdown('<div class="sb-dots">', unsafe_allow_html=True)
                        if st.button("⋮", key=f"pdots_{p[0]}"):
                            cur = st.session_state.get("open_menu")
                            st.session_state["open_menu"] = None if cur == f"player_{p[0]}" else f"player_{p[0]}"
                            st.rerun()
                        st.markdown('</div>', unsafe_allow_html=True)
                    if st.session_state.get("open_menu") == f"player_{p[0]}":
                        st.markdown('<div style="background:rgba(255,255,255,0.04);border:1px solid rgba(255,255,255,0.09);border-radius:8px;padding:6px 8px;margin:2px 14px 4px;">', unsafe_allow_html=True)
                        st.markdown('<div class="dot-edit">', unsafe_allow_html=True)
                        if st.button("✏  Edit Player", key=f"edit_p_{p[0]}"):
                            st.session_state["open_menu"] = None
                            st.session_state.show_edit_player = True
                            st.session_state.edit_player_id = p[0]
                            st.session_state.selected_player_id = p[0]
                            st.rerun()
                        st.markdown('</div><div class="dot-del">', unsafe_allow_html=True)
                        if st.button("✕  Delete Player", key=f"del_p_{p[0]}"):
                            st.session_state["open_menu"] = None
                            st.session_state["confirm_del_player_sb"] = p[0]
                            st.rerun()
                        st.markdown('</div></div>', unsafe_allow_html=True)
                    if st.session_state.get("confirm_del_player_sb") == p[0]:
                        st.markdown(f'<div style="background:rgba(239,68,68,0.08);border:1px solid rgba(239,68,68,0.2);border-radius:6px;padding:7px 12px;margin:2px 14px;font-size:11px;color:#fca5a5;">Delete {p[1]}?</div>', unsafe_allow_html=True)
                        dc1, dc2 = st.columns(2)
                        with dc1:
                            if st.button("Yes", key=f"cdp_{p[0]}"):
                                cursor.execute("DELETE FROM sessions WHERE player_id=?", (p[0],))
                                cursor.execute("DELETE FROM reports WHERE player_id=?", (p[0],))
                                cursor.execute("DELETE FROM players WHERE id=?", (p[0],))
                                conn.commit()
                                st.session_state["confirm_del_player_sb"] = None
                                st.session_state.selected_player_id = None
                                st.rerun()
                        with dc2:
                            if st.button("No", key=f"ndp_{p[0]}"):
                                st.session_state["confirm_del_player_sb"] = None; st.rerun()
                st.markdown('<div class="add-player-row">', unsafe_allow_html=True)
                if st.button("  ＋  Add Player", key=f"addp_{club}", use_container_width=True):
                    st.session_state.show_add_player = True
                    st.session_state.show_add_academy = False
                    st.session_state.add_player_club = club
                    st.session_state.selected_player_id = None
                    st.rerun()
                st.markdown('</div>', unsafe_allow_html=True)
    else:
        srch_e = st.text_input("se", placeholder="Search players...", label_visibility="collapsed")
        st.markdown('<div class="sb-section">Pro Teams <span class="badge-pro">EPL</span></div>', unsafe_allow_html=True)
        if epl_teams:
            # Show all pro players directly - team shown as subtitle
            cursor.execute("SELECT id,name,position,team FROM epl_players ORDER BY name")
            all_epl = cursor.fetchall()
            filt_e = [p for p in all_epl if srch_e.lower() in p[1].lower() or srch_e.lower() in (p[3] or "").lower()] if srch_e else all_epl
            if filt_e:
                # Group by team for display
                teams_seen = {}
                for p in filt_e:
                    t = p[3] or "No Team"
                    if t not in teams_seen: teams_seen[t] = []
                    teams_seen[t].append(p)
                for team, tplayers in teams_seen.items():
                    exp_e = team in st.session_state.expanded_epl_teams or bool(srch_e) or len(teams_seen) == 1
                    st.markdown('<div class="club-row">', unsafe_allow_html=True)
                    if st.button(f"{'▾' if exp_e else '▸'}  {team[:22]}", key=f"et_{team}", use_container_width=True):
                        if exp_e and not srch_e: st.session_state.expanded_epl_teams.remove(team) if team in st.session_state.expanded_epl_teams else None
                        elif not exp_e: st.session_state.expanded_epl_teams.append(team)
                        st.rerun()
                    st.markdown('</div>', unsafe_allow_html=True)
                    if exp_e:
                        for p in tplayers:
                            active = st.session_state.selected_epl_player_id == p[0]
                            st.markdown(f'<div class="{"player-active" if active else "player-row"}">', unsafe_allow_html=True)
                            if st.button(f"{'▸  ' if active else '   '}{p[1][:20]}", key=f"ep_{p[0]}", use_container_width=True):
                                st.session_state.selected_epl_player_id = p[0]; st.rerun()
                            st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.markdown('<div style="padding:10px 20px;font-size:11px;color:rgba(255,255,255,0.3);">Upload data in Pro Data tab</div>', unsafe_allow_html=True)
    st.markdown('<div class="sb-divider"></div>', unsafe_allow_html=True)
    st.markdown(f'<div class="sb-meta">Scout IQ·FC v1.0 Beta<br>{len(players)} Players · {len(epl_teams)} Pro Teams</div>', unsafe_allow_html=True)
    # Account + Sign Out buttons at bottom of sidebar
    st.markdown('<div style="padding:4px 16px 4px;">', unsafe_allow_html=True)
    if st.button("⚙  Account & Settings", key="open_account", use_container_width=True):
        st.session_state.show_account = not st.session_state.get("show_account", False)
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)
    st.markdown('<div class="signout-wrap">', unsafe_allow_html=True)
    if st.button("Sign Out", key="signout", use_container_width=True):
        st.session_state["authenticated"] = False
        st.session_state["auth_error"] = False
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

# ══════════════════════════════════════
# TOP RIGHT ACCOUNT BUTTON
# ══════════════════════════════════════
account_type = st.session_state.get("account_type", "admin")
badge_color = "#f5c842" if account_type == "admin" else "#3b82f6"
badge_label = "Admin" if account_type == "admin" else "Beta"
# Badge CSS only - no HTML injection that could fail before variables exist
# Account button - inject via sidebar bottom
# (Top-level columns break layout, use sidebar button instead)
# ══════════════════════════════════════
# ACCOUNT / SETTINGS PAGE
# ══════════════════════════════════════
if st.session_state.get("show_account"):
    st.markdown("""
    <style>
    .acct-page { max-width: 680px; margin: 0 auto; padding: 20px 0; }
    .acct-header { margin-bottom: 28px; }
    .acct-title { font-family: 'Playfair Display', serif; font-size: 32px; font-weight: 700; color: #04080f; margin-bottom: 4px; }
    .acct-sub { font-size: 12px; color: #9ca3af; letter-spacing: 1px; }
    .acct-card { background: #ffffff; border: 1px solid #e5e7ef; border-radius: 14px; padding: 24px 28px; margin-bottom: 16px; }
    .acct-card-title { font-size: 13px; font-weight: 700; color: #04080f; margin-bottom: 16px; display: flex; align-items: center; gap: 8px; }
    .acct-row { display: flex; justify-content: space-between; align-items: center; padding: 10px 0; border-bottom: 1px solid #f3f4f6; }
    .acct-row:last-child { border-bottom: none; }
    .acct-label { font-size: 12px; color: #6b7280; }
    .acct-value { font-size: 13px; font-weight: 600; color: #04080f; }
    .plan-badge { display: inline-block; background: #eff6ff; border: 1px solid #bfdbfe; color: #1d4ed8; font-size: 10px; font-weight: 700; letter-spacing: 1px; text-transform: uppercase; padding: 3px 10px; border-radius: 100px; }
    .plan-badge-beta { background: #fefce8; border-color: #fde68a; color: #92400e; }
    .danger-zone { border-color: #fecaca; }
    .danger-zone .acct-card-title { color: #dc2626; }
    </style>
    """, unsafe_allow_html=True)
    # Back button
    if st.button("← Back to Platform", key="close_account"):
        st.session_state.show_account = False
        st.rerun()
    st.markdown('<div class="acct-page">', unsafe_allow_html=True)
    # Page tabs
    acct_tab1, acct_tab2, acct_tab3 = st.tabs(["Account Overview", "Settings", "Billing & Plan"])
    with acct_tab1:
        is_admin = account_type == "admin"
        username = "admin" if is_admin else "admin1"
        st.markdown(f"""
        <div class="acct-card">
            <div class="acct-card-title">👤 Account Information</div>
            <div class="acct-row"><span class="acct-label">Username</span><span class="acct-value">{username}</span></div>
            <div class="acct-row"><span class="acct-label">Account Type</span><span class="acct-value">{"Administrator" if is_admin else "Beta User"}</span></div>
            <div class="acct-row"><span class="acct-label">Plan</span><span><span class="plan-badge {"" if is_admin else "plan-badge-beta"}">{"Pro — Admin" if is_admin else "Beta Access"}</span></span></div>
            <div class="acct-row"><span class="acct-label">Status</span><span class="acct-value" style="color:#16a34a;">● Active</span></div>
            <div class="acct-row"><span class="acct-label">Platform</span><span class="acct-value">Scout IQ·FC v1.0</span></div>
            <div class="acct-row"><span class="acct-label">Data Access</span><span class="acct-value">{"All Academies" if is_admin else "Riverside FC Academy only"}</span></div>
        </div>
        """, unsafe_allow_html=True)
        st.markdown(f"""
        <div class="acct-card">
            <div class="acct-card-title">📊 Usage Summary</div>
            <div class="acct-row"><span class="acct-label">Players in platform</span><span class="acct-value">{len(players)}</span></div>
            <div class="acct-row"><span class="acct-label">Sessions tracked</span><span class="acct-value">{cursor.execute("SELECT COUNT(*) FROM sessions").fetchone()[0]}</span></div>
            <div class="acct-row"><span class="acct-label">Reports generated</span><span class="acct-value">{cursor.execute("SELECT COUNT(*) FROM reports").fetchone()[0]}</span></div>
            <div class="acct-row"><span class="acct-label">Academies</span><span class="acct-value">{len(clubs)}</span></div>
            <div class="acct-row"><span class="acct-label">Pro Data players</span><span class="acct-value">{cursor.execute("SELECT COUNT(*) FROM epl_players").fetchone()[0]}</span></div>
        </div>
        """, unsafe_allow_html=True)
    with acct_tab2:
        st.markdown('<div class="acct-card"><div class="acct-card-title">🔐 Change Password</div>', unsafe_allow_html=True)
        current_pw = st.text_input("Current Password", type="password", key="chg_current")
        new_pw = st.text_input("New Password", type="password", key="chg_new")
        confirm_pw = st.text_input("Confirm New Password", type="password", key="chg_confirm")
        if st.button("Update Password", key="update_pw"):
            if new_pw and new_pw == confirm_pw:
                st.success("Password updated successfully. (Note: update your Streamlit secrets to persist this change.)")
            elif new_pw != confirm_pw:
                st.error("Passwords do not match.")
            else:
                st.error("Please fill in all fields.")
        st.markdown('</div>', unsafe_allow_html=True)
        st.markdown('<div class="acct-card"><div class="acct-card-title">🔔 Notifications</div>', unsafe_allow_html=True)
        st.toggle("Email me when a report is generated", value=True, key="notif_report")
        st.toggle("Weekly platform summary", value=False, key="notif_weekly")
        st.toggle("New feature announcements", value=True, key="notif_features")
        st.markdown('</div>', unsafe_allow_html=True)
        st.markdown('<div class="acct-card"><div class="acct-card-title">🎨 Preferences</div>', unsafe_allow_html=True)
        st.selectbox("Default report language", ["English", "Spanish", "French", "Arabic", "Portuguese"], key="pref_lang")
        st.selectbox("Default age group filter", ["All", "U13", "U14", "U15", "U16", "U17", "U18"], key="pref_age")
        st.markdown('</div>', unsafe_allow_html=True)
    with acct_tab3:
        is_admin_acct = account_type == "admin"
        plan_name = "Pro Admin" if is_admin_acct else "Beta Access"
        plan_price = "£49/month" if is_admin_acct else "Free (Beta)"
        plan_color = "#1d4ed8" if is_admin_acct else "#92400e"
        st.markdown(f"""
        <div class="acct-card" style="border-top:4px solid {plan_color};">
            <div class="acct-card-title">💳 Current Plan</div>
            <div class="acct-row"><span class="acct-label">Plan Name</span><span class="acct-value">{plan_name}</span></div>
            <div class="acct-row"><span class="acct-label">Price</span><span class="acct-value">{plan_price}</span></div>
            <div class="acct-row"><span class="acct-label">Billing Cycle</span><span class="acct-value">Monthly</span></div>
            <div class="acct-row"><span class="acct-label">Next Renewal</span><span class="acct-value">May 7, 2026</span></div>
            <div class="acct-row"><span class="acct-label">Status</span><span class="acct-value" style="color:#16a34a;">● Active</span></div>
        </div>
        """, unsafe_allow_html=True)
        st.markdown("""
        <div class="acct-card">
            <div class="acct-card-title">📦 Plan Features</div>
            <div class="acct-row"><span class="acct-label">Youth Players</span><span class="acct-value">Unlimited</span></div>
            <div class="acct-row"><span class="acct-label">AI Scouting Reports</span><span class="acct-value">Unlimited</span></div>
            <div class="acct-row"><span class="acct-label">Academies</span><span class="acct-value">Unlimited</span></div>
            <div class="acct-row"><span class="acct-label">Pro Data Import</span><span class="acct-value">✓ Included</span></div>
            <div class="acct-row"><span class="acct-label">PDF & Word Export</span><span class="acct-value">✓ Included</span></div>
            <div class="acct-row"><span class="acct-label">AI Smart Import</span><span class="acct-value">✓ Included</span></div>
        </div>
        """, unsafe_allow_html=True)
        st.markdown("""
        <div class="acct-card">
            <div class="acct-card-title">💳 Payment Method</div>
        """, unsafe_allow_html=True)
        st.markdown('<p style="font-size:12px;color:#9ca3af;margin-bottom:16px;">Payment processing powered by Stripe — coming soon</p>', unsafe_allow_html=True)
        pm1, pm2 = st.columns(2)
        with pm1:
            st.text_input("Card Number", placeholder="•••• •••• •••• ••••", key="card_num", disabled=True)
        with pm2:
            st.text_input("Expiry", placeholder="MM/YY", key="card_exp", disabled=True)
        if st.button("Update Payment Method", key="update_payment"):
            st.info("Stripe integration coming in the next update. Contact scoutiqfc@gmail.com to update billing.")
        st.markdown('</div>', unsafe_allow_html=True)
        # Available plans
        st.markdown('<div class="section-title" style="margin-top:24px;">Available Plans</div>', unsafe_allow_html=True)
        plan_col1, plan_col2, plan_col3 = st.columns(3)
        with plan_col1:
            st.markdown("""
            <div style="background:#f9fafb;border:1.5px solid #e5e7ef;border-radius:12px;padding:20px;text-align:center;">
                <div style="font-size:11px;font-weight:700;color:#6b7280;letter-spacing:2px;text-transform:uppercase;margin-bottom:8px;">Starter</div>
                <div style="font-size:28px;font-weight:800;color:#04080f;margin-bottom:4px;">£19</div>
                <div style="font-size:11px;color:#9ca3af;margin-bottom:16px;">per month</div>
                <div style="font-size:12px;color:#374151;line-height:1.8;">Up to 20 players<br>1 Academy<br>5 AI reports/month<br>PDF Export</div>
            </div>
            """, unsafe_allow_html=True)
        with plan_col2:
            st.markdown("""
            <div style="background:#eff6ff;border:2px solid #3b82f6;border-radius:12px;padding:20px;text-align:center;position:relative;">
                <div style="position:absolute;top:-10px;left:50%;transform:translateX(-50%);background:#3b82f6;color:#fff;font-size:9px;font-weight:700;letter-spacing:1px;padding:3px 10px;border-radius:100px;">POPULAR</div>
                <div style="font-size:11px;font-weight:700;color:#1d4ed8;letter-spacing:2px;text-transform:uppercase;margin-bottom:8px;">Pro</div>
                <div style="font-size:28px;font-weight:800;color:#04080f;margin-bottom:4px;">£49</div>
                <div style="font-size:11px;color:#9ca3af;margin-bottom:16px;">per month</div>
                <div style="font-size:12px;color:#374151;line-height:1.8;">Unlimited players<br>Unlimited academies<br>Unlimited AI reports<br>Pro Data Import<br>AI Smart Import</div>
            </div>
            """, unsafe_allow_html=True)
        with plan_col3:
            st.markdown("""
            <div style="background:#fafaf9;border:1.5px solid #e5e7ef;border-radius:12px;padding:20px;text-align:center;">
                <div style="font-size:11px;font-weight:700;color:#6b7280;letter-spacing:2px;text-transform:uppercase;margin-bottom:8px;">Academy</div>
                <div style="font-size:28px;font-weight:800;color:#04080f;margin-bottom:4px;">£149</div>
                <div style="font-size:11px;color:#9ca3af;margin-bottom:16px;">per month</div>
                <div style="font-size:12px;color:#374151;line-height:1.8;">5 coach accounts<br>White label reports<br>Custom branding<br>Priority support<br>API access</div>
            </div>
            """, unsafe_allow_html=True)
        st.markdown('<p style="font-size:12px;color:#9ca3af;margin-top:16px;text-align:center;">To upgrade or change your plan, contact <a href="mailto:scoutiqfc@gmail.com" style="color:#3b82f6;">scoutiqfc@gmail.com</a></p>', unsafe_allow_html=True)
        # Danger zone
        st.markdown("""
        <div class="acct-card danger-zone" style="margin-top:24px;border-color:#fecaca;">
            <div class="acct-card-title" style="color:#dc2626;">⚠ Danger Zone</div>
        """, unsafe_allow_html=True)
        if st.button("Cancel Subscription", key="cancel_sub"):
            st.warning("To cancel your subscription contact scoutiqfc@gmail.com. Your data will be retained for 30 days.")
        if st.button("Delete Account", key="delete_acct"):
            st.error("Account deletion is permanent. Contact scoutiqfc@gmail.com to proceed.")
        st.markdown('</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)
    st.stop()
# ══════════════════════════════════════
# EDIT ACADEMY PAGE
# ══════════════════════════════════════
if st.session_state.get("show_edit_academy") and st.session_state.get("edit_academy_name"):
    club = st.session_state.edit_academy_name
    cursor.execute("SELECT * FROM players WHERE club=?", (club,))
    academy_players = cursor.fetchall()
    if st.button("← Back", key="back_edit_acad"):
        st.session_state.show_edit_academy = False
        st.rerun()
    st.markdown(f'<div class="page-header"><div><div class="page-title">{club}</div><div class="page-meta">Academy Management</div></div></div>', unsafe_allow_html=True)
    ea1, ea2 = st.tabs(["Academy Details", "Player Roster"])
    with ea1:
        st.markdown('<div class="section-title">Edit Academy</div>', unsafe_allow_html=True)
        new_name = st.text_input("Academy Name", value=club, key="edit_acad_name_input")
        about = st.text_area("About this Academy", placeholder="Founded year, location, philosophy, youth development approach...", height=100, key="edit_acad_about")
        contact = st.text_input("Contact Email", placeholder="academy@club.com", key="edit_acad_contact")
        if st.button("Save Changes", key="save_acad_edit"):
            if new_name and new_name.strip() != club:
                cursor.execute("UPDATE players SET club=? WHERE club=?", (new_name.strip(), club))
                conn.commit()
                st.session_state.edit_academy_name = new_name.strip()
                st.success(f"Academy renamed to {new_name}")
                st.rerun()
            else:
                st.success("Academy details saved.")
    with ea2:
        st.markdown('<div class="section-title">Player Roster</div>', unsafe_allow_html=True)
        if academy_players:
            for ap in academy_players:
                apc1, apc2, apc3 = st.columns([5, 1, 1])
                with apc1:
                    st.markdown(f'<div style="padding:8px 0;font-size:13px;color:#374151;font-weight:500;">{ap[1]} <span style="font-size:11px;color:#9ca3af;">· {ap[4]} · {ap[3]}</span></div>', unsafe_allow_html=True)
                with apc2:
                    if st.button("View", key=f"view_ap_{ap[0]}"):
                        st.session_state.selected_player_id = ap[0]
                        st.session_state.show_edit_academy = False
                        st.rerun()
                with apc3:
                    st.markdown('<div class="sb-action-del">', unsafe_allow_html=True)
                    if st.button("✕", key=f"del_ap_{ap[0]}"):
                        cursor.execute("DELETE FROM sessions WHERE player_id=?", (ap[0],))
                        cursor.execute("DELETE FROM reports WHERE player_id=?", (ap[0],))
                        cursor.execute("DELETE FROM players WHERE id=?", (ap[0],))
                        conn.commit()
                        st.rerun()
                    st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.info("No players in this academy yet.")
    st.stop()
# ══════════════════════════════════════
# YOUTH MODE
# ══════════════════════════════════════
if st.session_state.mode == "youth":
    # CREATE ACADEMY
    if st.session_state.get("show_add_academy"):
        st.markdown('<div class="page-header"><div><div class="page-title">New Academy</div><div class="page-meta">Create a squad or club</div></div></div>', unsafe_allow_html=True)
        academy_name = st.text_input("Academy name", placeholder="e.g. Riverside FC Academy", key="acad_name")
        c1, c2 = st.columns([1,5])
        with c1:
            if st.button("Create", key="create_acad"):
                if academy_name and academy_name.strip():
                    st.session_state.show_add_academy = False
                    st.session_state.show_add_player = True
                    st.session_state.add_player_club = academy_name.strip()
                    st.rerun()
                else:
                    st.error("Enter an academy name")
        with c2:
            if st.button("Cancel", key="cancel_acad"):
                st.session_state.show_add_academy = False; st.rerun()
        st.stop()
    # ADD PLAYER
    if st.session_state.get("show_add_player"):
        club_name = st.session_state.get("add_player_club", "Unknown Club")
        st.markdown(f'<div class="page-header"><div><div class="page-title">Add Player</div><div class="page-meta">{club_name}</div></div></div>', unsafe_allow_html=True)
        tab_manual, tab_upload, tab_ai = st.tabs(["Manual Entry", "Upload Excel", "✨ AI Smart Import"])
        with tab_manual:
            with st.form("add_player_form"):
                c1,c2,c3 = st.columns(3)
                with c1:
                    p_name = st.text_input("Full Name *")
                    p_pos = st.selectbox("Position", ["Striker","Right Winger","Left Winger","Attacking Midfielder","Central Midfielder","Defensive Midfielder","Right Back","Left Back","Centre Back","Goalkeeper"])
                    p_dob = st.text_input("Date of Birth", placeholder="YYYY-MM-DD")
                with c2:
                    p_age = st.selectbox("Age Group", ["U13","U14","U15","U16","U17","U18"])
                    p_foot = st.selectbox("Dominant Foot", ["Right","Left"])
                    p_nat = st.text_input("Nationality")
                with c3:
                    s_date = st.text_input("Session Date", placeholder="YYYY-MM-DD")
                    s_type = st.selectbox("Session Type", ["match","training"])
                    s_mins = st.number_input("Minutes", 0, 120, 90)
                st.markdown("**Session Metrics**")
                sc1,sc2,sc3,sc4 = st.columns(4)
                with sc1:
                    s_dist=st.number_input("Distance km",0.0,15.0,8.0,0.1)
                    s_spr=st.number_input("Sprints",0,60,12)
                    s_spd=st.number_input("Top Speed",0.0,40.0,28.0,0.1)
                with sc2:
                    s_pcomp=st.number_input("Passes Completed",0,120,30)
                    s_patt=st.number_input("Passes Attempted",0,120,38)
                    s_drib=st.number_input("Dribbles",0,30,4)
                with sc3:
                    s_def=st.number_input("Defensive Actions",0,40,6)
                    s_goals=st.number_input("Goals",0,10,0)
                    s_ast=st.number_input("Assists",0,10,0)
                with sc4:
                    s_tack=st.number_input("Tackles Won",0,25,3)
                    s_coach=st.slider("Coachability",1,10,7)
                    s_att_s=st.slider("Attitude",1,10,7)
                    s_cons=st.slider("Consistency",1,10,7)
                s_notes=st.text_area("Coach Notes")
                if st.form_submit_button("Add Player"):
                    if not p_name or not p_name.strip():
                        st.error("Player name required")
                    else:
                        cursor.execute("SELECT id FROM players WHERE name=?", (p_name.strip(),))
                        ex = cursor.fetchone()
                        if ex:
                            new_pid = ex[0]
                        else:
                            cursor.execute("INSERT INTO players (name,date_of_birth,position,club,dominant_foot,age_group,nationality) VALUES (?,?,?,?,?,?,?)",
                                (p_name.strip(),p_dob,p_pos,club_name,p_foot,p_age,p_nat))
                            conn.commit(); new_pid = cursor.lastrowid
                        cursor.execute("""INSERT INTO sessions (player_id,session_date,session_type,minutes_played,distance_covered_km,sprint_count,top_speed_kmh,passes_completed,passes_attempted,dribbles_completed,defensive_actions,goals,assists,chances_created,tackles_won,coachability_rating,attitude_score,consistency_rating,coach_notes) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                            (new_pid,s_date,s_type,s_mins,s_dist,s_spr,s_spd,s_pcomp,s_patt,s_drib,s_def,s_goals,s_ast,0,s_tack,s_coach,s_att_s,s_cons,s_notes))
                        conn.commit()
                        st.session_state.selected_player_id = new_pid
                        st.session_state.show_add_player = False
                        st.success(f"{p_name} added.")
                        st.rerun()
        with tab_upload:
            st.markdown('<div class="upload-banner"><p style="font-size:13px;font-weight:600;color:#1d4ed8;margin-bottom:4px;">Upload Excel to import players</p><p style="font-size:12px;color:#6b7280;margin:0;">Use Players.xlsx column format. Each row is one session.</p></div>', unsafe_allow_html=True)
            up_file = st.file_uploader("Upload", type=["xlsx"], label_visibility="collapsed")
            if up_file:
                try:
                    df_up = pd.read_excel(up_file)
                    st.dataframe(df_up.head(3), use_container_width=True)
                    if st.button("Import All", key="imp_all"):
                        imported = 0
                        for _, row in df_up.iterrows():
                            try:
                                pname = str(row.get("name","")).strip()
                                if not pname or pname=="nan": continue
                                cursor.execute("SELECT id FROM players WHERE name=?", (pname,))
                                ex = cursor.fetchone()
                                if ex: new_pid = ex[0]
                                else:
                                    cursor.execute("INSERT INTO players (name,date_of_birth,position,club,dominant_foot,age_group,nationality) VALUES (?,?,?,?,?,?,?)",
                                        (pname,str(row.get("date_of_birth",""))[:10],str(row.get("position","")),club_name,str(row.get("dominant_foot","Right")),str(row.get("age_group","U16")),str(row.get("nationality",""))))
                                    conn.commit(); new_pid = cursor.lastrowid
                                def gv(c, d=0, fl=False):
                                    try:
                                        v=row.get(c,d)
                                        if pd.isna(v): return d
                                        return float(v) if fl else int(float(v))
                                    except: return d
                                cursor.execute("""INSERT INTO sessions (player_id,session_date,session_type,minutes_played,distance_covered_km,sprint_count,top_speed_kmh,passes_completed,passes_attempted,dribbles_completed,defensive_actions,goals,assists,chances_created,tackles_won,coachability_rating,attitude_score,consistency_rating,coach_notes) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                                    (new_pid,str(row.get("session_date",""))[:10],str(row.get("session_type","match")),gv("minutes_played"),gv("distance_covered_km",0,True),gv("sprint_count"),gv("top_speed_kmh",0,True),gv("passes_completed"),gv("passes_attempted"),gv("dribbles_completed"),gv("defensive_actions"),gv("goals"),gv("assists"),gv("chances_created"),gv("tackles_won"),gv("coachability_rating"),gv("attitude_score"),gv("consistency_rating"),str(row.get("coach_notes",""))))
                                conn.commit(); imported += 1
                            except: continue
                        st.success(f"Imported {imported} records")
                        st.session_state.show_add_player = False; st.rerun()
                except Exception as e:
                    st.error(f"Error: {e}")
        with tab_ai:
            st.markdown('<div style="background:#eff6ff;border:1px solid #bfdbfe;border-radius:12px;padding:16px 20px;margin-bottom:16px;"><p style="font-size:13px;font-weight:700;color:#1d4ed8;margin-bottom:4px;">✨ AI Smart Import — Any Format</p><p style="font-size:12px;color:#3b82f6;margin:0;">Upload a file OR chat below. Drop in match notes, WhatsApp messages, photos of scoresheets, PDFs, Word docs, Excel files, or just type anything.</p></div>', unsafe_allow_html=True)
            # File upload
            ai_file = st.file_uploader("Upload file (photo, PDF, Word, Excel, CSV)", type=["jpg","jpeg","png","pdf","docx","xlsx","csv","txt"], key=f"youth_ai_file_{club_name}", label_visibility="collapsed")
            # Chat/text input
            raw_input = st.text_area("Or describe / paste data here",
                placeholder="Type naturally e.g Marcus 15 striker scored 2 goals. Or paste any table CSV match report.",
                height=130, key="ai_raw_input", label_visibility="collapsed")
            # Chat history for context
            if "ai_chat_history" not in st.session_state:
                st.session_state.ai_chat_history = []
            if st.button("✨ Parse with AI", key="ai_parse_btn"):
                input_text = raw_input.strip() if raw_input else ""
                with st.spinner("Claude is reading your data..."):
                    try:
                        import json, io as _io
                        if ai_file:
                            file_bytes = ai_file.read()
                            ext = ai_file.name.split(".")[-1].lower()
                            if ext in ["jpg","jpeg","png","webp"]:
                                import base64
                                b64 = base64.standard_b64encode(file_bytes).decode()
                                media_map = {"jpg":"image/jpeg","jpeg":"image/jpeg","png":"image/png","webp":"image/webp"}
                                claude_client = anthropic.Anthropic(api_key=get_api_key())
                                r = claude_client.messages.create(
                                    model="claude-opus-4-5", max_tokens=3000,
                                    messages=[{"role":"user","content":[
                                        {"type":"image","source":{"type":"base64","media_type":media_map.get(ext,"image/jpeg"),"data":b64}},
                                        {"type":"text","text":"Extract all youth football player data from this image. Return a JSON array where each object has: name, position, goals, assists, minutes_played, coach_notes. Return ONLY the JSON array, no markdown."}
                                    ]}])
                                raw_result = r.content[0].text
                            elif ext == "docx":
                                from docx import Document as _DocxDoc
                                try:
                                    _d = _DocxDoc(io.BytesIO(file_bytes))
                                    input_text = "\n".join([p.text for p in _d.paragraphs if p.text.strip()])
                                except Exception:
                                    input_text = file_bytes.decode("utf-8","ignore")[:6000]
                            elif ext in ["xlsx","xls"]:
                                _df = pd.read_excel(io.BytesIO(file_bytes))
                                raw_result = ai_clean_data(_df.to_string(index=False)[:5000], "youth")
                            elif ext in ["csv","txt"]:
                                input_text = file_bytes.decode("utf-8","ignore")[:6000]
                                raw_result = ai_clean_data(input_text, "youth")
                            else:
                                input_text = file_bytes.decode("utf-8","ignore")[:6000]
                                raw_result = ai_clean_data(input_text, "youth")
                        elif input_text:
                            raw_result = ai_clean_data(input_text, "youth")
                        else:
                            st.warning("Upload a file or type some data first."); st.stop()
                        raw_result = raw_result.strip()
                        if raw_result.startswith("```"):
                            parts = raw_result.split("```")
                            raw_result = parts[1] if len(parts)>1 else raw_result
                            if raw_result.startswith("json"): raw_result = raw_result[4:]
                        players_data = json.loads(raw_result.strip())
                        if not isinstance(players_data, list): players_data = [players_data]
                        st.session_state["ai_parsed_players"] = players_data
                        st.session_state["ai_parse_club"] = club_name
                    except Exception as e:
                        st.error(f"Could not parse: {e}. Try adding more detail or structure.")
            if st.session_state.get("ai_parsed_players"):
                players_data = st.session_state["ai_parsed_players"]
                st.success(f"Claude found {len(players_data)} player record(s) — review before saving:")
                prev_df = pd.DataFrame([{"Name":p.get("name","?"),"Position":p.get("position","?"),"Goals":p.get("goals",0),"Assists":p.get("assists",0),"Coach":p.get("coachability_rating",7)} for p in players_data])
                st.dataframe(prev_df, use_container_width=True, hide_index=True)
                sc1, sc2 = st.columns([1,4])
                with sc1:
                    if st.button("✅ Save All to Platform", key="ai_confirm_save"):
                        saved = 0
                        for p in players_data:
                            pname = str(p.get("name","")).strip()
                            if not pname: continue
                            p_club = p.get("club", club_name) or club_name
                            cursor.execute("SELECT id FROM players WHERE name=?", (pname,))
                            ex = cursor.fetchone()
                            if ex: new_pid = ex[0]
                            else:
                                cursor.execute("INSERT INTO players (name,date_of_birth,position,club,dominant_foot,age_group,nationality) VALUES (?,?,?,?,?,?,?)",(pname,p.get("date_of_birth",""),p.get("position","Central Midfielder"),p_club,p.get("dominant_foot","Right"),p.get("age_group","U16"),p.get("nationality","")))
                                conn.commit(); new_pid = cursor.lastrowid
                            def sv(key, default=0, fl=False):
                                try:
                                    v=p.get(key,default); return float(v) if fl else int(float(v)) if v is not None else default
                                except: return default
                            cursor.execute("INSERT INTO sessions (player_id,session_date,session_type,minutes_played,distance_covered_km,sprint_count,top_speed_kmh,passes_completed,passes_attempted,dribbles_completed,defensive_actions,goals,assists,chances_created,tackles_won,coachability_rating,attitude_score,consistency_rating,coach_notes) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",(new_pid,p.get("session_date","2024-09-01"),p.get("session_type","match"),sv("minutes_played",90),sv("distance_covered_km",8,True),sv("sprint_count",15),sv("top_speed_kmh",28,True),sv("passes_completed",25),sv("passes_attempted",32),sv("dribbles_completed",4),sv("defensive_actions",6),sv("goals"),sv("assists"),sv("chances_created"),sv("tackles_won"),sv("coachability_rating",7),sv("attitude_score",7),sv("consistency_rating",7),str(p.get("coach_notes",""))))
                            conn.commit(); saved += 1
                        st.session_state["ai_parsed_players"] = None
                        st.session_state.selected_player_id = new_pid
                        st.session_state.show_add_player = False
                        st.success(f"{saved} player(s) saved. Opening player profile...")
                        st.rerun()
                with sc2:
                    if st.button("✕ Discard", key="ai_discard"):
                        st.session_state["ai_parsed_players"] = None; st.rerun()
        if st.button("← Cancel", key="cancel_add"):
            st.session_state.show_add_player = False; st.rerun()
        st.stop()
    # ── ACADEMY LANDING PAGE ──
    if st.session_state.get("selected_academy") and not st.session_state.selected_player_id:
        club = st.session_state.selected_academy
        acad_ps = clubs.get(club, [])
        if st.button("← Back", key="bk_acad_lp"):
            st.session_state.selected_academy = None; st.rerun()
        if not acad_ps:
            st.markdown(f'<div class="page-header"><div class="page-title">{club}</div></div>', unsafe_allow_html=True)
            st.info("No players in this academy yet.")
            if st.button("＋ Add First Player", key="add_first_p"):
                st.session_state.show_add_player = True
                st.session_state.add_player_club = club
                st.session_state.selected_academy = None
                st.rerun()
            st.stop()
        all_pids = [p[0] for p in acad_ps]
        ph = ",".join(["?"]*len(all_pids))
        r = cursor.execute(f"SELECT SUM(goals),SUM(assists),SUM(minutes_played),COUNT(*) FROM sessions WHERE player_id IN ({ph})", all_pids).fetchone()
        tot_g=r[0] or 0; tot_a=r[1] or 0; tot_mins=r[2] or 0; tot_sess=r[3] or 0
        r2 = cursor.execute(f"SELECT AVG(coachability_rating),AVG(attitude_score),AVG(consistency_rating),AVG(distance_covered_km),AVG(sprint_count) FROM sessions WHERE player_id IN ({ph})", all_pids).fetchone()
        avg_coach=round(r2[0] or 0,1); avg_att=round(r2[1] or 0,1); avg_cons=round(r2[2] or 0,1)
        avg_dist=round(r2[3] or 0,1); avg_spr=round(r2[4] or 0,1)
        tot_reps = cursor.execute(f"SELECT COUNT(*) FROM reports WHERE player_id IN ({ph})", all_pids).fetchone()[0]
        # Hero
        st.markdown(f"""
        <div style="background:linear-gradient(135deg,#04080f 0%,#1a3a8a 55%,#0f1f5c 100%);border-radius:16px;padding:36px 40px;margin-bottom:24px;position:relative;overflow:hidden;">
            <div style="position:absolute;inset:0;background-image:linear-gradient(rgba(255,255,255,0.02) 1px,transparent 1px),linear-gradient(90deg,rgba(255,255,255,0.02) 1px,transparent 1px);background-size:44px 44px;"></div>
            <div style="position:relative;z-index:1;display:flex;justify-content:space-between;align-items:flex-start;">
                <div>
                    <div style="font-size:9px;font-weight:700;color:rgba(245,200,66,0.65);letter-spacing:4px;text-transform:uppercase;margin-bottom:10px;">Scout IQ·FC · Academy</div>
                    <div style="font-family:'Playfair Display',serif;font-size:38px;font-weight:700;color:#fff;letter-spacing:-0.5px;margin-bottom:8px;">{club}</div>
                    <div style="font-size:12px;color:rgba(255,255,255,0.4);">{len(acad_ps)} players · {tot_sess} sessions · {tot_reps} reports generated</div>
                </div>
                <div style="text-align:right;">
                    <div style="font-size:9px;font-weight:700;color:rgba(245,200,66,0.6);letter-spacing:2px;text-transform:uppercase;margin-bottom:4px;">Season Goals</div>
                    <div style="font-size:52px;font-weight:800;color:#f5c842;line-height:1;">{int(tot_g)}</div>
                </div>
            </div>
            <div style="display:flex;gap:36px;margin-top:24px;position:relative;z-index:1;flex-wrap:wrap;">
                <div style="text-align:center;"><div style="font-size:22px;font-weight:700;color:#fff;">{int(tot_a)}</div><div style="font-size:9px;color:rgba(255,255,255,0.32);letter-spacing:2px;text-transform:uppercase;margin-top:2px;">Assists</div></div>
                <div style="text-align:center;"><div style="font-size:22px;font-weight:700;color:#fff;">{avg_coach}/10</div><div style="font-size:9px;color:rgba(255,255,255,0.32);letter-spacing:2px;text-transform:uppercase;margin-top:2px;">Avg Coach</div></div>
                <div style="text-align:center;"><div style="font-size:22px;font-weight:700;color:#fff;">{avg_att}/10</div><div style="font-size:9px;color:rgba(255,255,255,0.32);letter-spacing:2px;text-transform:uppercase;margin-top:2px;">Avg Attitude</div></div>
                <div style="text-align:center;"><div style="font-size:22px;font-weight:700;color:#fff;">{avg_cons}/10</div><div style="font-size:9px;color:rgba(255,255,255,0.32);letter-spacing:2px;text-transform:uppercase;margin-top:2px;">Consistency</div></div>
                <div style="text-align:center;"><div style="font-size:22px;font-weight:700;color:#fff;">{avg_dist}km</div><div style="font-size:9px;color:rgba(255,255,255,0.32);letter-spacing:2px;text-transform:uppercase;margin-top:2px;">Avg Distance</div></div>
                <div style="text-align:center;"><div style="font-size:22px;font-weight:700;color:#fff;">{avg_spr}</div><div style="font-size:9px;color:rgba(255,255,255,0.32);letter-spacing:2px;text-transform:uppercase;margin-top:2px;">Avg Sprints</div></div>
            </div>
        </div>
        """, unsafe_allow_html=True)
        # Stats row
        mc1,mc2,mc3,mc4,mc5 = st.columns(5)
        for col,lbl,val,tp in [(mc1,"Players",len(acad_ps),"#1a3a8a"),(mc2,"Sessions",tot_sess,"#3b82f6"),(mc3,"Goals",int(tot_g),"#f5c842"),(mc4,"Assists",int(tot_a),"#16a34a"),(mc5,"Reports",tot_reps,"#7c3aed")]:
            with col:
                st.markdown(f'<div class="metric-card" style="border-top:3px solid {tp};"><div class="metric-label">{lbl}</div><div class="metric-value">{val}</div></div>', unsafe_allow_html=True)
        # Charts row
        import plotly.graph_objects as go
        pos_counts = {}
        age_counts = {}
        for p in acad_ps:
            pos_counts[p[2] or "Unknown"] = pos_counts.get(p[2] or "Unknown", 0) + 1
            age_counts[p[4] or "Unknown"] = age_counts.get(p[4] or "Unknown", 0) + 1
        st.markdown('<div class="section-title">Squad Analytics</div>', unsafe_allow_html=True)
        ch1,ch2,ch3 = st.columns(3)
        with ch1:
            fig = go.Figure(go.Bar(x=list(pos_counts.keys()), y=list(pos_counts.values()), marker_color='#1a3a8a', opacity=0.85, marker_line_width=0))
            fig.update_layout(title='Squad by Position', paper_bgcolor='white', plot_bgcolor='white', margin=dict(l=10,r=10,t=36,b=60), xaxis=dict(showgrid=False, tickangle=45), yaxis=dict(gridcolor='#f3f4f6'), height=230, showlegend=False)
            st.plotly_chart(fig, use_container_width=True)
        with ch2:
            fig = go.Figure(go.Bar(x=list(age_counts.keys()), y=list(age_counts.values()), marker_color='#f5c842', opacity=0.9, marker_line_width=0))
            fig.update_layout(title='Squad by Age Group', paper_bgcolor='white', plot_bgcolor='white', margin=dict(l=10,r=10,t=36,b=20), xaxis=dict(showgrid=False), yaxis=dict(gridcolor='#f3f4f6'), height=230, showlegend=False)
            st.plotly_chart(fig, use_container_width=True)
        with ch3:
            player_goals = []
            for p in acad_ps:
                g = cursor.execute("SELECT SUM(goals) FROM sessions WHERE player_id=?", (p[0],)).fetchone()[0] or 0
                if g > 0: player_goals.append((p[1][:14], g))
            player_goals.sort(key=lambda x: x[1], reverse=True)
            if player_goals:
                fig = go.Figure(go.Bar(x=[x[0] for x in player_goals[:8]], y=[x[1] for x in player_goals[:8]], marker_color='#16a34a', opacity=0.85, marker_line_width=0))
                fig.update_layout(title='Top Scorers', paper_bgcolor='white', plot_bgcolor='white', margin=dict(l=10,r=10,t=36,b=60), xaxis=dict(showgrid=False, tickangle=45), yaxis=dict(gridcolor='#f3f4f6'), height=230, showlegend=False)
                st.plotly_chart(fig, use_container_width=True)
        # Player roster
        # Academy actions row
        aa1, aa2, aa3 = st.columns([3, 1, 1])
        with aa1:
            st.markdown('<div class="section-title">Player Roster</div>', unsafe_allow_html=True)
        with aa2:
            st.markdown('<div style="padding-top:24px;">', unsafe_allow_html=True)
            if st.button("＋ Add Player", key="add_from_acad_lp", use_container_width=True):
                st.session_state.show_add_player = True
                st.session_state.add_player_club = club
                st.session_state.selected_academy = None
                st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)
        with aa3:
            st.markdown('<div style="padding-top:24px;">', unsafe_allow_html=True)
            if st.button("✨ AI Import", key="ai_import_acad", use_container_width=True):
                st.session_state.show_ai_import_academy = club
                st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)
        # AI Import panel for academy
        if st.session_state.get("show_ai_import_academy") == club:
            with st.expander("✨ AI Smart Import — Upload or paste any data", expanded=True):
                st.markdown('<div style="background:#eff6ff;border:1px solid #bfdbfe;border-radius:8px;padding:10px 14px;margin-bottom:12px;font-size:12px;color:#1d4ed8;">Upload a photo, PDF, Word doc, Excel, or just type notes. Claude will extract all player data automatically.</div>', unsafe_allow_html=True)
                ai_acad_file = st.file_uploader("Upload file", type=["jpg","jpeg","png","pdf","docx","xlsx","csv","txt"], key=f"ai_acad_file_{club}", label_visibility="collapsed")
                ai_acad_text = st.text_area("Or paste / type any player data", placeholder="Type naturally e.g: Marcus 15 striker scored 2 goals, great attitude. Or paste any table or match notes.", height=100, key=f"ai_acad_text_{club}", label_visibility="collapsed")
                parse_col, close_col = st.columns([1, 4])
                with parse_col:
                    parse_btn = st.button("✨ Parse", key=f"ai_acad_parse_{club}")
                with close_col:
                    if st.button("✕ Close", key=f"ai_acad_close_{club}"):
                        st.session_state.show_ai_import_academy = None; st.rerun()
                if parse_btn:
                    input_text = ai_acad_text.strip() if ai_acad_text else ""
                    with st.spinner("Claude is reading your data..."):
                        try:
                            import json
                            if ai_acad_file:
                                file_bytes = ai_acad_file.read()
                                ext = ai_acad_file.name.split(".")[-1].lower()
                                if ext in ["jpg","jpeg","png","webp"]:
                                    b64 = base64.standard_b64encode(file_bytes).decode()
                                    media_map = {"jpg":"image/jpeg","jpeg":"image/jpeg","png":"image/png","webp":"image/webp"}
                                    cl = anthropic.Anthropic(api_key=get_api_key())
                                    r = cl.messages.create(model="claude-opus-4-5", max_tokens=2000, messages=[{"role":"user","content":[{"type":"image","source":{"type":"base64","media_type":media_map.get(ext,"image/jpeg"),"data":b64}},{"type":"text","text":"Extract all youth football player data from this image. Return a JSON array where each object has: name, position, goals, assists, minutes_played, coach_notes. Return ONLY valid JSON, no markdown."}]}])
                                    raw_result = r.content[0].text
                                elif ext == "docx":
                                    from docx import Document as _D
                                    try:
                                        input_text = "\n".join([p.text for p in _D(io.BytesIO(file_bytes)).paragraphs if p.text.strip()])
                                    except:
                                        input_text = file_bytes.decode("utf-8","ignore")[:6000]
                                elif ext in ["xlsx","xls"]:
                                    raw_result = ai_clean_data(pd.read_excel(io.BytesIO(file_bytes)).to_string(index=False)[:5000], "youth")
                                else:
                                    raw_result = ai_clean_data(file_bytes.decode("utf-8","ignore")[:6000], "youth")
                            elif input_text:
                                raw_result = ai_clean_data(input_text, "youth")
                            else:
                                st.warning("Upload a file or type some data first."); st.stop()
                            raw_result = raw_result.strip()
                            if raw_result.startswith("```"):
                                parts = raw_result.split("```")
                                raw_result = parts[1][4:] if len(parts) > 1 and parts[1].startswith("json") else (parts[1] if len(parts) > 1 else raw_result)
                            parsed = json.loads(raw_result.strip())
                            if not isinstance(parsed, list): parsed = [parsed]
                            st.session_state[f"ai_acad_parsed_{club}"] = parsed
                        except Exception as e:
                            st.error(f"Could not parse: {str(e)[:100]}")
                if st.session_state.get(f"ai_acad_parsed_{club}"):
                    parsed = st.session_state[f"ai_acad_parsed_{club}"]
                    st.success(f"Found {len(parsed)} player(s) — confirm to save:")
                    st.dataframe(pd.DataFrame([{"Name": p.get("name","?"), "Position": p.get("position","?"), "Goals": p.get("goals",0), "Assists": p.get("assists",0)} for p in parsed]), use_container_width=True, hide_index=True)
                    sv1, sv2 = st.columns([1,4])
                    with sv1:
                        if st.button("✅ Save All", key=f"ai_acad_save_{club}"):
                            saved = 0
                            for p in parsed:
                                pname = str(p.get("name","")).strip()
                                if not pname: continue
                                cursor.execute("SELECT id FROM players WHERE name=?", (pname,))
                                ex = cursor.fetchone()
                                if ex: npid = ex[0]
                                else:
                                    cursor.execute("INSERT INTO players (name,date_of_birth,position,club,dominant_foot,age_group,nationality) VALUES (?,?,?,?,?,?,?)", (pname,p.get("date_of_birth",""),p.get("position","Central Midfielder"),club,p.get("dominant_foot","Right"),p.get("age_group","U16"),p.get("nationality","")))
                                    conn.commit(); npid = cursor.lastrowid
                                def sv_val(k,d=0,fl=False):
                                    try:
                                        v=p.get(k,d); return float(v) if fl else int(float(v)) if v is not None else d
                                    except: return d
                                cursor.execute("INSERT INTO sessions (player_id,session_date,session_type,minutes_played,distance_covered_km,sprint_count,top_speed_kmh,passes_completed,passes_attempted,dribbles_completed,defensive_actions,goals,assists,chances_created,tackles_won,coachability_rating,attitude_score,consistency_rating,coach_notes) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)", (npid,p.get("session_date","2024-09-01"),p.get("session_type","match"),sv_val("minutes_played",90),sv_val("distance_covered_km",8,True),sv_val("sprint_count",15),sv_val("top_speed_kmh",28,True),sv_val("passes_completed",25),sv_val("passes_attempted",32),sv_val("dribbles_completed",4),sv_val("defensive_actions",6),sv_val("goals"),sv_val("assists"),sv_val("chances_created"),sv_val("tackles_won"),sv_val("coachability_rating",7),sv_val("attitude_score",7),sv_val("consistency_rating",7),str(p.get("coach_notes",""))))
                                conn.commit(); saved += 1
                            st.session_state[f"ai_acad_parsed_{club}"] = None
                            st.session_state.show_ai_import_academy = None
                            st.success(f"{saved} player(s) added to {club}"); st.rerun()
                    with sv2:
                        if st.button("✕ Discard", key=f"ai_acad_discard_{club}"):
                            st.session_state[f"ai_acad_parsed_{club}"] = None; st.rerun()
        for i in range(0, len(acad_ps), 3):
            cols = st.columns(3)
            for j, p in enumerate(acad_ps[i:i+3]):
                with cols[j]:
                    ps = cursor.execute("SELECT COUNT(*),SUM(goals),SUM(assists),AVG(coachability_rating) FROM sessions WHERE player_id=?", (p[0],)).fetchone()
                    p_sess=ps[0] or 0; p_g=int(ps[1] or 0); p_a=int(ps[2] or 0); p_c=round(ps[3] or 0,1)
                    has_rep = cursor.execute("SELECT COUNT(*) FROM reports WHERE player_id=?", (p[0],)).fetchone()[0] > 0
                    badge = '<span style="background:#eff6ff;border:1px solid #bfdbfe;color:#1d4ed8;font-size:8px;font-weight:700;padding:2px 6px;border-radius:100px;margin-left:6px;">✓ Report</span>' if has_rep else ''
                    st.markdown(f"""
                    <div style="background:#fff;border:1px solid #e5e7ef;border-radius:12px;padding:16px 18px;margin-bottom:8px;">
                        <div style="font-size:14px;font-weight:700;color:#04080f;margin-bottom:3px;">{p[1]}{badge}</div>
                        <div style="font-size:11px;color:#9ca3af;margin-bottom:10px;">{p[2]} · {p[4]} · {p[5] or ''}</div>
                        <div style="display:flex;gap:14px;">
                            <div><div style="font-size:9px;color:#9ca3af;text-transform:uppercase;letter-spacing:1px;">Sessions</div><div style="font-size:13px;font-weight:700;color:#374151;">{p_sess}</div></div>
                            <div><div style="font-size:9px;color:#9ca3af;text-transform:uppercase;letter-spacing:1px;">Goals</div><div style="font-size:13px;font-weight:700;color:#374151;">{p_g}</div></div>
                            <div><div style="font-size:9px;color:#9ca3af;text-transform:uppercase;letter-spacing:1px;">Assists</div><div style="font-size:13px;font-weight:700;color:#374151;">{p_a}</div></div>
                            <div><div style="font-size:9px;color:#9ca3af;text-transform:uppercase;letter-spacing:1px;">Coach</div><div style="font-size:13px;font-weight:700;color:#374151;">{p_c}/10</div></div>
                        </div>
                    </div>
                    """, unsafe_allow_html=True)
                    if st.button(f"View Profile →", key=f"view_lp_{p[0]}"):
                        st.session_state.selected_player_id = p[0]
                        st.session_state.selected_academy = None
                        st.rerun()
        st.stop()

    # ── ACADEMY LANDING PAGE ──
    if st.session_state.get("selected_academy") and not st.session_state.selected_player_id:
        club = st.session_state.selected_academy
        acad_ps = clubs.get(club, [])
        if st.button("← Back", key="bk_acad_lp"):
            st.session_state.selected_academy = None; st.rerun()
        if not acad_ps:
            st.markdown(f'<div class="page-header"><div class="page-title">{club}</div></div>', unsafe_allow_html=True)
            st.info("No players yet.")
            st.stop()
        all_pids = [p[0] for p in acad_ps]
        ph = ",".join(["?"]*len(all_pids))
        r = cursor.execute(f"SELECT SUM(goals),SUM(assists),COUNT(*) FROM sessions WHERE player_id IN ({ph})", all_pids).fetchone()
        tot_g=r[0] or 0; tot_a=r[1] or 0; tot_sess=r[2] or 0
        r2 = cursor.execute(f"SELECT AVG(coachability_rating),AVG(attitude_score),AVG(consistency_rating),AVG(distance_covered_km),AVG(sprint_count) FROM sessions WHERE player_id IN ({ph})", all_pids).fetchone()
        avg_coach=round(r2[0] or 0,1); avg_att=round(r2[1] or 0,1); avg_cons=round(r2[2] or 0,1); avg_dist=round(r2[3] or 0,1); avg_spr=round(r2[4] or 0,1)
        tot_reps = cursor.execute(f"SELECT COUNT(*) FROM reports WHERE player_id IN ({ph})", all_pids).fetchone()[0]
        st.markdown(f"""<div style="background:linear-gradient(135deg,#04080f 0%,#1a3a8a 55%,#0f1f5c 100%);border-radius:16px;padding:36px 40px;margin-bottom:24px;">
<div style="font-size:9px;font-weight:700;color:rgba(245,200,66,0.65);letter-spacing:4px;text-transform:uppercase;margin-bottom:10px;">Scout IQ·FC · Academy</div>
<div style="font-family:'Playfair Display',serif;font-size:38px;font-weight:700;color:#fff;letter-spacing:-0.5px;margin-bottom:8px;">{club}</div>
<div style="font-size:12px;color:rgba(255,255,255,0.4);">{len(acad_ps)} players · {tot_sess} sessions · {tot_reps} reports</div>
<div style="display:flex;gap:32px;margin-top:20px;flex-wrap:wrap;">
<div style="text-align:center;"><div style="font-size:22px;font-weight:700;color:#f5c842;">{int(tot_g)}</div><div style="font-size:9px;color:rgba(255,255,255,0.32);letter-spacing:2px;text-transform:uppercase;margin-top:2px;">Goals</div></div>
<div style="text-align:center;"><div style="font-size:22px;font-weight:700;color:#fff;">{int(tot_a)}</div><div style="font-size:9px;color:rgba(255,255,255,0.32);letter-spacing:2px;text-transform:uppercase;margin-top:2px;">Assists</div></div>
<div style="text-align:center;"><div style="font-size:22px;font-weight:700;color:#fff;">{avg_coach}/10</div><div style="font-size:9px;color:rgba(255,255,255,0.32);letter-spacing:2px;text-transform:uppercase;margin-top:2px;">Avg Coach</div></div>
<div style="text-align:center;"><div style="font-size:22px;font-weight:700;color:#fff;">{avg_att}/10</div><div style="font-size:9px;color:rgba(255,255,255,0.32);letter-spacing:2px;text-transform:uppercase;margin-top:2px;">Avg Attitude</div></div>
<div style="text-align:center;"><div style="font-size:22px;font-weight:700;color:#fff;">{avg_dist}km</div><div style="font-size:9px;color:rgba(255,255,255,0.32);letter-spacing:2px;text-transform:uppercase;margin-top:2px;">Avg Dist</div></div>
</div></div>""", unsafe_allow_html=True)
        mc1,mc2,mc3,mc4 = st.columns(4)
        for col,lbl,val,tp in [(mc1,"Players",len(acad_ps),"#1a3a8a"),(mc2,"Sessions",tot_sess,"#3b82f6"),(mc3,"Goals",int(tot_g),"#f5c842"),(mc4,"Reports",tot_reps,"#16a34a")]:
            with col:
                st.markdown(f'<div class="metric-card" style="border-top:3px solid {tp};"><div class="metric-card-label">{lbl}</div><div class="metric-card-value">{val}</div></div>', unsafe_allow_html=True)
        rc1, rc2 = st.columns([5,1])
        with rc1: st.markdown('<div class="section-title">Player Roster</div>', unsafe_allow_html=True)
        with rc2:
            if st.button("＋ Add Player", key="add_from_acad"):
                st.session_state.show_add_player = True
                st.session_state.add_player_club = club
                st.session_state.selected_academy = None
                st.rerun()
        for i in range(0, len(acad_ps), 3):
            cols = st.columns(3)
            for j, p in enumerate(acad_ps[i:i+3]):
                with cols[j]:
                    ps = cursor.execute("SELECT COUNT(*),SUM(goals),SUM(assists),AVG(coachability_rating) FROM sessions WHERE player_id=?", (p[0],)).fetchone()
                    p_sess=ps[0] or 0; p_g=int(ps[1] or 0); p_a=int(ps[2] or 0); p_c=round(ps[3] or 0,1)
                    st.markdown(f"""<div style="background:#fff;border:1px solid #e5e7ef;border-radius:12px;padding:16px 18px;margin-bottom:8px;">
<div style="font-size:14px;font-weight:700;color:#04080f;margin-bottom:3px;">{p[1]}</div>
<div style="font-size:11px;color:#9ca3af;margin-bottom:10px;">{p[2]} · {p[4]}</div>
<div style="display:flex;gap:14px;">
<div><div style="font-size:9px;color:#9ca3af;text-transform:uppercase;letter-spacing:1px;">Sessions</div><div style="font-size:13px;font-weight:700;color:#374151;">{p_sess}</div></div>
<div><div style="font-size:9px;color:#9ca3af;text-transform:uppercase;letter-spacing:1px;">Goals</div><div style="font-size:13px;font-weight:700;color:#374151;">{p_g}</div></div>
<div><div style="font-size:9px;color:#9ca3af;text-transform:uppercase;letter-spacing:1px;">Assists</div><div style="font-size:13px;font-weight:700;color:#374151;">{p_a}</div></div>
<div><div style="font-size:9px;color:#9ca3af;text-transform:uppercase;letter-spacing:1px;">Coach</div><div style="font-size:13px;font-weight:700;color:#374151;">{p_c}/10</div></div>
</div></div>""", unsafe_allow_html=True)
                    if st.button(f"View →", key=f"view_lp_{p[0]}"):
                        st.session_state.selected_player_id = p[0]
                        st.session_state.selected_academy = None
                        st.rerun()
        st.stop()

    # WELCOME SCREEN
    if not st.session_state.selected_player_id:
        st.markdown('<div class="page-header"><div><div class="page-title">Scout IQ·FC</div><div class="page-meta">Youth Talent Intelligence Platform</div></div></div>', unsafe_allow_html=True)
        c1,c2,c3 = st.columns(3)
        with c1:
            st.markdown('<div class="welcome-card"><div class="welcome-icon">🏟️</div><div class="welcome-title">Create an Academy</div><div class="welcome-desc">Click New Academy in the sidebar to register your club</div></div>', unsafe_allow_html=True)
        with c2:
            st.markdown('<div class="welcome-card"><div class="welcome-icon">👤</div><div class="welcome-title">Add Players</div><div class="welcome-desc">Expand any academy and click Add Player to build your roster</div></div>', unsafe_allow_html=True)
        with c3:
            st.markdown('<div class="welcome-card"><div class="welcome-icon">📋</div><div class="welcome-title">Generate Reports</div><div class="welcome-desc">Select a player and generate a full AI scouting report</div></div>', unsafe_allow_html=True)
        st.stop()
    # EDIT PLAYER PAGE
    if st.session_state.get("show_edit_player") and st.session_state.get("edit_player_id"):
        epid_edit = st.session_state.edit_player_id
        cursor.execute("SELECT * FROM players WHERE id=?", (epid_edit,))
        ep_edit = cursor.fetchone()
        if ep_edit:
            if st.button("← Back to Player", key="back_edit_player"):
                st.session_state.show_edit_player = False
                st.rerun()
            st.markdown(f'<div class="page-header"><div><div class="page-title">{ep_edit[1]}</div><div class="page-meta">Edit Player Profile</div></div></div>', unsafe_allow_html=True)
            ep_tab1, ep_tab2, ep_tab3 = st.tabs(["Profile", "Sessions", "Reports"])
            with ep_tab1:
                st.markdown('<div class="section-title">Player Details</div>', unsafe_allow_html=True)
                with st.form("edit_player_form"):
                    ec1, ec2, ec3 = st.columns(3)
                    with ec1:
                        new_pname = st.text_input("Full Name", value=ep_edit[1])
                        new_pos = st.selectbox("Position", ["Striker","Right Winger","Left Winger","Attacking Midfielder","Central Midfielder","Defensive Midfielder","Right Back","Left Back","Centre Back","Goalkeeper"],
                            index=["Striker","Right Winger","Left Winger","Attacking Midfielder","Central Midfielder","Defensive Midfielder","Right Back","Left Back","Centre Back","Goalkeeper"].index(ep_edit[4]) if ep_edit[4] in ["Striker","Right Winger","Left Winger","Attacking Midfielder","Central Midfielder","Defensive Midfielder","Right Back","Left Back","Centre Back","Goalkeeper"] else 0)
                        new_dob = st.text_input("Date of Birth", value=ep_edit[2] or "")
                    with ec2:
                        new_age = st.selectbox("Age Group", ["U13","U14","U15","U16","U17","U18"],
                            index=["U13","U14","U15","U16","U17","U18"].index(ep_edit[3]) if ep_edit[3] in ["U13","U14","U15","U16","U17","U18"] else 0)
                        new_foot = st.selectbox("Dominant Foot", ["Right","Left"],
                            index=0 if ep_edit[5]=="Right" else 1)
                        new_nat = st.text_input("Nationality", value=ep_edit[7] or "")
                    with ec3:
                        new_club = st.text_input("Academy / Club", value=ep_edit[6] or "")
                    if st.form_submit_button("Save Changes"):
                        cursor.execute("UPDATE players SET name=?,date_of_birth=?,age_group=?,position=?,dominant_foot=?,club=?,nationality=? WHERE id=?",
                            (new_pname,new_dob,new_age,new_pos,new_foot,new_club,new_nat,epid_edit))
                        conn.commit()
                        st.success("Player profile updated.")
                        st.rerun()
            with ep_tab2:
                st.markdown('<div class="section-title">Session History</div>', unsafe_allow_html=True)
                cursor.execute("SELECT * FROM sessions WHERE player_id=? ORDER BY session_date DESC", (epid_edit,))
                edit_sessions = cursor.fetchall()
                if edit_sessions:
                    for s in edit_sessions:
                        sc1, sc2 = st.columns([8,1])
                        with sc1:
                            st.markdown(f'<div style="padding:8px 0;font-size:12px;color:#374151;">{str(s[2])[:10]} · {s[3].title()} · {s[4]} mins · {s[12]}G {s[13]}A</div>', unsafe_allow_html=True)
                        with sc2:
                            st.markdown('<div class="sb-action-del">', unsafe_allow_html=True)
                            if st.button("✕", key=f"del_sess_{s[0]}"):
                                cursor.execute("DELETE FROM sessions WHERE id=?", (s[0],))
                                conn.commit()
                                st.rerun()
                            st.markdown('</div>', unsafe_allow_html=True)
                else:
                    st.info("No sessions recorded yet.")
            with ep_tab3:
                st.markdown('<div class="section-title">Reports</div>', unsafe_allow_html=True)
                cursor.execute("SELECT * FROM reports WHERE player_id=? ORDER BY created_at DESC", (epid_edit,))
                edit_reports = cursor.fetchall()
                if edit_reports:
                    for r in edit_reports:
                        rc1, rc2 = st.columns([8,1])
                        with rc1:
                            st.markdown(f'<div style="padding:8px 0;font-size:12px;color:#374151;">Report generated {str(r[3])[:10]}</div>', unsafe_allow_html=True)
                        with rc2:
                            st.markdown('<div class="sb-action-del">', unsafe_allow_html=True)
                            if st.button("✕", key=f"del_rep_{r[0]}"):
                                cursor.execute("DELETE FROM reports WHERE id=?", (r[0],))
                                conn.commit()
                                st.rerun()
                            st.markdown('</div>', unsafe_allow_html=True)
                else:
                    st.info("No reports generated yet.")
            st.stop()
    # PLAYER DASHBOARD
    pid = st.session_state.selected_player_id
    cursor.execute("SELECT * FROM players WHERE id=?", (pid,))
    pl = cursor.fetchone()
    if not pl:
        st.session_state.selected_player_id = None; st.rerun()
    cursor.execute("SELECT * FROM sessions WHERE player_id=? ORDER BY session_date", (pid,))
    sess = cursor.fetchall()
    cursor.execute("SELECT * FROM reports WHERE player_id=? ORDER BY created_at DESC LIMIT 1", (pid,))
    rep = cursor.fetchone()
    sc = len(sess)
    avg_d = round(sum(s[5] for s in sess)/sc,1) if sc else 0
    avg_sp = round(sum(s[6] for s in sess)/sc,1) if sc else 0
    tg = sum(s[12] for s in sess); ta = sum(s[13] for s in sess)
    tpc = sum(s[8] for s in sess); tpa = sum(s[9] for s in sess)
    pp = round(tpc/tpa*100) if tpa>0 else 0
    # Header with generate button
    hc1, hc2 = st.columns([4,1])
    with hc1:
        st.markdown(f'<div class="page-header"><div><div class="page-title">{pl[1]}</div><div class="page-meta">{pl[4]} &nbsp;·&nbsp; {pl[6]} &nbsp;·&nbsp; {pl[3]} &nbsp;·&nbsp; {pl[7] or ""}</div></div></div>', unsafe_allow_html=True)
    # No generate button at top - single button at bottom only
    c1,c2,c3,c4,c5,c6 = st.columns(6)
    cards = [(c1,"Sessions",sc,"neutral"),(c2,"Avg Distance",f"{avg_d} km","neutral"),
             (c3,"Avg Sprints",avg_sp,"neutral"),(c4,"Pass Rate",f"{pp}%","highlight"),
             (c5,"Goals",tg,"positive"),(c6,"Assists",ta,"positive")]
    for col,lbl,val,style in cards:
        with col:
            st.markdown(f'<div class="metric-card {style}"><div class="metric-label">{lbl}</div><div class="metric-value">{val}</div></div>', unsafe_allow_html=True)
    # Session history
    # Session history header + add button - no overlap
    sh_c1, sh_c2 = st.columns([6, 1])
    with sh_c1:
        st.markdown('<div class="section-title" style="margin-bottom:0;">Session History</div>', unsafe_allow_html=True)
    with sh_c2:
        st.markdown('<div style="padding-top:28px;">', unsafe_allow_html=True)
        add_key = f"show_add_sess_{pid}"
        if add_key not in st.session_state: st.session_state[add_key] = False
        if st.button("＋ Add", key=f"add_s_{pid}"):
            st.session_state[add_key] = not st.session_state[add_key]; st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
    if st.session_state[add_key]:
        with st.expander("New Session", expanded=True):
            with st.form(f"new_sess_{pid}"):
                fc1,fc2,fc3 = st.columns(3)
                with fc1:
                    ns_date=st.text_input("Date",placeholder="YYYY-MM-DD",key=f"nsd_{pid}")
                    ns_type=st.selectbox("Type",["match","training"],key=f"nst_{pid}")
                    ns_mins=st.number_input("Minutes",0,120,90,key=f"nsm_{pid}")
                with fc2:
                    ns_dist=st.number_input("Distance km",0.0,15.0,8.0,0.1,key=f"nsdist_{pid}")
                    ns_spr=st.number_input("Sprints",0,60,12,key=f"nsspr_{pid}")
                    ns_spd=st.number_input("Top Speed",0.0,40.0,28.0,0.1,key=f"nsspd_{pid}")
                    ns_pc=st.number_input("Passes Comp",0,120,30,key=f"nspc_{pid}")
                    ns_pa=st.number_input("Passes Att",0,120,38,key=f"nspa_{pid}")
                with fc3:
                    ns_g=st.number_input("Goals",0,10,0,key=f"nsg_{pid}")
                    ns_a=st.number_input("Assists",0,10,0,key=f"nsa_{pid}")
                    ns_def=st.number_input("Def Actions",0,40,6,key=f"nsdef_{pid}")
                    ns_tack=st.number_input("Tackles",0,25,3,key=f"nstk_{pid}")
                    ns_coach=st.slider("Coachability",1,10,7,key=f"nsc_{pid}")
                    ns_att=st.slider("Attitude",1,10,7,key=f"nsat_{pid}")
                    ns_cons=st.slider("Consistency",1,10,7,key=f"nscon_{pid}")
                ns_notes=st.text_area("Notes",key=f"nsn_{pid}")
                ns_excel=st.file_uploader("Or upload Excel",type=["xlsx"],key=f"nse_{pid}",label_visibility="collapsed")
                if st.form_submit_button("Save Session"):
                    if ns_excel:
                        try:
                            df_ns=pd.read_excel(ns_excel); added=0
                            for _,row in df_ns.iterrows():
                                def gv2(c,d=0,fl=False):
                                    try:
                                        v=row.get(c,d)
                                        if pd.isna(v): return d
                                        return float(v) if fl else int(float(v))
                                    except: return d
                                cursor.execute("""INSERT INTO sessions (player_id,session_date,session_type,minutes_played,distance_covered_km,sprint_count,top_speed_kmh,passes_completed,passes_attempted,dribbles_completed,defensive_actions,goals,assists,chances_created,tackles_won,coachability_rating,attitude_score,consistency_rating,coach_notes) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                                    (pid,str(row.get("session_date",""))[:10],str(row.get("session_type","match")),gv2("minutes_played"),gv2("distance_covered_km",0,True),gv2("sprint_count"),gv2("top_speed_kmh",0,True),gv2("passes_completed"),gv2("passes_attempted"),gv2("dribbles_completed"),gv2("defensive_actions"),gv2("goals"),gv2("assists"),gv2("chances_created"),gv2("tackles_won"),gv2("coachability_rating"),gv2("attitude_score"),gv2("consistency_rating"),str(row.get("coach_notes",""))))
                                conn.commit(); added+=1
                            st.session_state[add_key]=False
                            st.success(f"Added {added} sessions"); st.rerun()
                        except Exception as e: st.error(f"Error: {e}")
                    elif ns_date:
                        cursor.execute("""INSERT INTO sessions (player_id,session_date,session_type,minutes_played,distance_covered_km,sprint_count,top_speed_kmh,passes_completed,passes_attempted,dribbles_completed,defensive_actions,goals,assists,chances_created,tackles_won,coachability_rating,attitude_score,consistency_rating,coach_notes) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                            (pid,ns_date,ns_type,ns_mins,ns_dist,ns_spr,ns_spd,ns_pc,ns_pa,0,ns_def,ns_g,ns_a,0,ns_tack,ns_coach,ns_att,ns_cons,ns_notes))
                        conn.commit(); st.session_state[add_key]=False
                        st.success("Session saved"); st.rerun()
                    else:
                        st.error("Enter a date or upload Excel")
    if sess:
        rows=[{"Date":str(s[2])[:10],"Type":s[3].title(),"Mins":s[4],"Dist":f"{s[5]}km","Sprints":s[6],"Speed":f"{s[7]}km/h","Pass%":f"{round(s[8]/s[9]*100) if s[9]>0 else 0}%","Goals":s[12],"Assists":s[13],"Tackles":s[15],"Coach":f"{s[16]}/10","Att":f"{s[17]}/10"} for s in sess]
        st.dataframe(pd.DataFrame(rows), use_container_width=True, hide_index=True)
        st.markdown('<div class="section-title">Performance Charts</div>', unsafe_allow_html=True)
        dates=[str(s[2])[:10] for s in sess]
        c1,c2=st.columns(2)
        with c1:
            fig=go.Figure()
            fig.add_trace(go.Scatter(x=dates,y=[s[5] for s in sess],mode='lines+markers',line=dict(color='#1a3a8a',width=2.5),marker=dict(size=7,color='#1a3a8a'),fill='tozeroy',fillcolor='rgba(26,58,138,0.06)'))
            fig.update_layout(title='Distance per Session',paper_bgcolor='white',plot_bgcolor='white',margin=dict(l=10,r=10,t=40,b=10),xaxis=dict(showgrid=False,type='category'),yaxis=dict(gridcolor='#f3f4f6',title='km'),height=240,showlegend=False)
            st.plotly_chart(fig,use_container_width=True)
        with c2:
            fig=go.Figure()
            fig.add_trace(go.Bar(x=dates,y=[s[6] for s in sess],marker_color='#3b82f6',opacity=0.8,marker_line_width=0))
            fig.update_layout(title='Sprint Count',paper_bgcolor='white',plot_bgcolor='white',margin=dict(l=10,r=10,t=40,b=10),xaxis=dict(showgrid=False,type='category'),yaxis=dict(gridcolor='#f3f4f6'),height=240,showlegend=False)
            st.plotly_chart(fig,use_container_width=True)
        c3,c4=st.columns(2)
        with c3:
            fig=go.Figure()
            fig.add_trace(go.Scatter(x=dates,y=[round(s[8]/s[9]*100) if s[9]>0 else 0 for s in sess],mode='lines+markers',line=dict(color='#16a34a',width=2.5),marker=dict(size=7,color='#16a34a'),fill='tozeroy',fillcolor='rgba(22,163,74,0.06)'))
            fig.update_layout(title='Pass Accuracy %',paper_bgcolor='white',plot_bgcolor='white',margin=dict(l=10,r=10,t=40,b=10),xaxis=dict(showgrid=False,type='category'),yaxis=dict(gridcolor='#f3f4f6',title='%',range=[0,100]),height=240,showlegend=False)
            st.plotly_chart(fig,use_container_width=True)
        with c4:
            cats=['Distance','Sprints','Passing','Attacking','Defending','Attitude']
            maxv=[12,25,100,10,20,10]
            avgv=[avg_d,avg_sp,pp,round((tg+ta+sum(s[14] for s in sess))/sc,1),round(sum(s[15] for s in sess)/sc,1),round(sum(s[17] for s in sess)/sc,1)]
            norm=[min(round((v/m)*10,1),10) for v,m in zip(avgv,maxv)]
            fig=go.Figure()
            fig.add_trace(go.Scatterpolar(r=norm+[norm[0]],theta=cats+[cats[0]],fill='toself',fillcolor='rgba(26,58,138,0.1)',line=dict(color='#1a3a8a',width=2),marker=dict(size=5)))
            fig.update_layout(title='Performance Profile',polar=dict(radialaxis=dict(visible=True,range=[0,10],gridcolor='#e5e7ef'),angularaxis=dict(gridcolor='#e5e7ef'),bgcolor='white'),paper_bgcolor='white',margin=dict(l=10,r=10,t=40,b=10),height=240,showlegend=False)
            st.plotly_chart(fig,use_container_width=True)
        st.markdown('<div class="section-title">Development Indicators</div>', unsafe_allow_html=True)
        fig=go.Figure()
        fig.add_trace(go.Scatter(x=dates,y=[s[16] for s in sess],name='Coachability',line=dict(color='#1a3a8a',width=2),mode='lines+markers',marker=dict(size=6)))
        fig.add_trace(go.Scatter(x=dates,y=[s[17] for s in sess],name='Attitude',line=dict(color='#f5c842',width=2,dash='dash'),mode='lines+markers',marker=dict(size=6)))
        fig.add_trace(go.Scatter(x=dates,y=[s[18] for s in sess],name='Consistency',line=dict(color='#16a34a',width=2,dash='dot'),mode='lines+markers',marker=dict(size=6)))
        fig.update_layout(title='Development Scores',paper_bgcolor='white',plot_bgcolor='white',margin=dict(l=10,r=10,t=40,b=10),xaxis=dict(showgrid=False,type='category'),yaxis=dict(gridcolor='#f3f4f6',range=[0,10],title='Score /10'),legend=dict(orientation='h',yanchor='bottom',y=1.02,xanchor='right',x=1),height=260)
        st.plotly_chart(fig,use_container_width=True)
    # Media
    st.markdown('<div class="section-title">Player Media</div>', unsafe_allow_html=True)
    pc,vc=st.columns(2)
    with pc:
        st.markdown('<p style="font-size:10px;font-weight:700;color:#9ca3af;letter-spacing:2px;text-transform:uppercase;margin-bottom:8px;">Player Photo</p>', unsafe_allow_html=True)
        photo=st.file_uploader("photo",type=["jpg","jpeg","png"],key=f"ph_{pid}",label_visibility="collapsed")
        if photo: st.image(photo,width=200)
        else: st.markdown('<div style="background:#f9fafb;border:2px dashed #e5e7ef;border-radius:12px;padding:24px;text-align:center;color:#9ca3af;font-size:12px;">Upload player photo</div>', unsafe_allow_html=True)
    with vc:
        st.markdown('<p style="font-size:10px;font-weight:700;color:#9ca3af;letter-spacing:2px;text-transform:uppercase;margin-bottom:8px;">Match Footage</p>', unsafe_allow_html=True)
        video=st.file_uploader("video",type=["mp4","mov","avi"],key=f"vi_{pid}",label_visibility="collapsed")
        if video: st.video(video)
        else: st.markdown('<div style="background:#f9fafb;border:2px dashed #e5e7ef;border-radius:12px;padding:24px;text-align:center;color:#9ca3af;font-size:12px;">Upload match footage</div>', unsafe_allow_html=True)
    # Scouting Report
    st.markdown('<div class="section-title">Scouting Report</div>', unsafe_allow_html=True)


    if rep:
        render_report(rep[2], pid, pl[1], f"{pl[4]}  ·  {pl[6]}  ·  {pl[3]}", is_pro=False)
    else:
        st.markdown('<div style="background:#f9fafb;border:1px solid #e5e7ef;border-radius:16px;padding:48px;text-align:center;">', unsafe_allow_html=True)
        st.markdown('<p style="font-size:32px;margin-bottom:12px;">📋</p><p style="font-size:15px;font-weight:600;color:#374151;margin-bottom:6px;">No report generated yet</p><p style="font-size:13px;color:#9ca3af;margin-bottom:24px;">Generate a full AI scouting report based on session data</p>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        st.markdown('<div class="gen-btn" style="text-align:center;margin-top:16px;">', unsafe_allow_html=True)
        if st.button("Generate AI Scouting Report", key=f"gen_bottom_{pid}"):
            cursor.execute("SELECT * FROM sessions WHERE player_id=? ORDER BY session_date", (pid,))
            fresh_sess = cursor.fetchall()
            if not fresh_sess:
                st.warning("Add session data first.")
            else:
                with st.spinner("Generating report — about 30 seconds..."):
                    rtext = ai_report(build_youth_prompt(pl, fresh_sess))
                    cursor.execute("INSERT INTO reports (player_id,report_text) VALUES (?,?)", (pid, rtext))
                    conn.commit(); st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

# ══════════════════════════════════════
# PRO DATA MODE
# ══════════════════════════════════════
else:
    st.markdown('<div class="page-header"><div><div class="page-title">Pro Data</div><div class="page-meta">Professional Match Analysis &nbsp;·&nbsp; Real Performance Data</div></div></div>', unsafe_allow_html=True)
    tab1,tab2,tab3,tab4=st.tabs(["Upload Data","Player Report","Team Analysis","✨ AI Import"])
    with tab1:
        st.markdown('<div class="section-title">Upload Professional Data</div>', unsafe_allow_html=True)
        st.markdown('<div class="upload-banner"><p style="font-size:13px;font-weight:600;color:#1d4ed8;margin-bottom:4px;">Import Professional Player Data</p><p style="font-size:12px;color:#6b7280;margin:0;">Supports season stats files (FootyStats, FBref, WhoScored) or gameweek-by-gameweek match data. Your Premierleague_Data.xlsx format is fully supported.</p></div>', unsafe_allow_html=True)
        uploaded=st.file_uploader("Upload",type=["xlsx","csv"],label_visibility="collapsed")
        if uploaded:
            try:
                df=pd.read_csv(uploaded) if uploaded.name.endswith('.csv') else pd.read_excel(uploaded)
                st.success(f"{len(df)} rows loaded — {df.shape[1]} columns detected")
                st.dataframe(df.head(3),use_container_width=True)
                # Auto-detect format
                cols_lower = [str(c).lower() for c in df.columns]
                is_season_format = any(c in cols_lower for c in ['full_name','goals_overall','minutes_played_overall','appearances_overall'])
                is_gameweek_format = any(c in cols_lower for c in ['gameweek','opponent','xg','xa'])
                if is_season_format:
                    st.info("Season stats format detected — this will import each player as one record with their full season totals.")
                    if st.button("Import All Players", key="imp_season"):
                        imported = 0
                        errors = []
                        def gv(row, col, default=0, fl=False):
                            try:
                                v = row.get(col, default)
                                if pd.isna(v): return default
                                return float(v) if fl else int(float(v))
                            except: return default
                        for _, row in df.iterrows():
                            try:
                                # Get player name - try full_name first, then name
                                pname = str(row.get("full_name", row.get("name", ""))).strip()
                                if not pname or pname == "nan": continue
                                team = str(row.get("Current Club", row.get("team", row.get("club", "Unknown")))).strip()
                                position = str(row.get("position", "")).strip()
                                nationality = str(row.get("nationality", "")).strip()
                                age = gv(row, "age")
                                season = str(row.get("season", "2024/25")).strip()
                                # Check if player already exists
                                cursor.execute("SELECT id FROM epl_players WHERE name=? AND team=?", (pname, team))
                                ex = cursor.fetchone()
                                if ex:
                                    epid = ex[0]
                                else:
                                    cursor.execute("INSERT INTO epl_players (name,team,position,nationality,age) VALUES (?,?,?,?,?)",
                                        (pname, team, position, nationality, age))
                                    conn.commit()
                                    epid = cursor.lastrowid
                                # Map season stats to a single session record
                                mins = gv(row, "minutes_played_overall")
                                goals = gv(row, "goals_overall")
                                assists = gv(row, "assists_overall")
                                appearances = gv(row, "appearances_overall")
                                yellow_cards = gv(row, "yellow_cards_overall")
                                red_cards = gv(row, "red_cards_overall")
                                goals_p90 = gv(row, "goals_per_90_overall", 0, True)
                                assists_p90 = gv(row, "assists_per_90_overall", 0, True)
                                clean_sheets = gv(row, "clean_sheets_overall")
                                # Store as gameweek=0 meaning full season
                                cursor.execute("""INSERT INTO epl_sessions
                                    (player_id, gameweek, opponent, minutes_played, goals, assists,
                                    shots, shots_on_target, xg, xa, passes_completed, tackles_won,
                                    yellow_cards, red_cards, rating)
                                    VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                                    (epid, 0, f"Full Season {season}",
                                    mins, goals, assists,
                                    0, 0,
                                    round(goals_p90 * mins/90, 2) if mins > 0 else 0,
                                    round(assists_p90 * mins/90, 2) if mins > 0 else 0,
                                    appearances, clean_sheets,
                                    yellow_cards, red_cards,
                                    round(goals_p90 + assists_p90, 2)))
                                conn.commit()
                                imported += 1
                            except Exception as e:
                                errors.append(f"{pname}: {e}")
                                continue
                        if imported > 0:
                            st.success(f"Successfully imported {imported} players")
                            if errors:
                                st.warning(f"{len(errors)} rows skipped")
                            st.rerun()
                        else:
                            st.error("No players imported. Check the file format.")
                            for e in errors[:5]: st.write(e)
                else:
                    # Manual column mapping for gameweek format
                    team_name=st.text_input("Team name",placeholder="e.g. Manchester United")
                    cols=["(skip)"]+list(df.columns)
                    with st.expander("Map columns"):
                        cc1,cc2,cc3=st.columns(3)
                        with cc1:
                            nc=st.selectbox("Player name",cols,key="nc"); poc=st.selectbox("Position",cols,key="poc")
                            agc=st.selectbox("Age",cols,key="agc"); natc=st.selectbox("Nationality",cols,key="natc")
                        with cc2:
                            mc=st.selectbox("Minutes",cols,key="mc"); gc=st.selectbox("Goals",cols,key="gc")
                            ac=st.selectbox("Assists",cols,key="ac"); shc=st.selectbox("Shots",cols,key="shc")
                            sotc=st.selectbox("Shots on target",cols,key="sotc")
                        with cc3:
                            xgc=st.selectbox("xG",cols,key="xgc"); xac=st.selectbox("xA",cols,key="xac")
                            pcc=st.selectbox("Passes completed",cols,key="pcc"); tc=st.selectbox("Tackles",cols,key="tc")
                            gwc=st.selectbox("Gameweek",cols,key="gwc"); oppc=st.selectbox("Opponent",cols,key="oppc")
                            ratc=st.selectbox("Rating",cols,key="ratc")
                    if st.button("Import Data",key="imp"):
                        if not team_name: st.error("Enter team name")
                        elif nc=="(skip)": st.error("Map player name column")
                        else:
                            def sv(row,col,default=0,fl=False):
                                if col=="(skip)": return default
                                try:
                                    v=row.get(col,default)
                                    if pd.isna(v): return default
                                    return float(v) if fl else int(float(v))
                                except: return default
                            imported=0
                            for _,row in df.iterrows():
                                pname=str(row.get(nc,"")).strip()
                                if not pname or pname=="nan": continue
                                cursor.execute("SELECT id FROM epl_players WHERE name=? AND team=?",(pname,team_name))
                                ex=cursor.fetchone()
                                if ex: epid=ex[0]
                                else:
                                    cursor.execute("INSERT INTO epl_players (name,team,position,nationality,age) VALUES (?,?,?,?,?)",(pname,team_name,str(row.get(poc,"")) if poc!="(skip)" else "",str(row.get(natc,"")) if natc!="(skip)" else "",sv(row,agc)))
                                    conn.commit(); epid=cursor.lastrowid
                                cursor.execute("INSERT INTO epl_sessions (player_id,gameweek,opponent,minutes_played,goals,assists,shots,shots_on_target,xg,xa,passes_completed,tackles_won,rating) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)",
                                    (epid,sv(row,gwc),str(row.get(oppc,"")) if oppc!="(skip)" else "",sv(row,mc),sv(row,gc),sv(row,ac),sv(row,shc),sv(row,sotc),sv(row,xgc,0,True),sv(row,xac,0,True),sv(row,pcc),sv(row,tc),sv(row,ratc,0,True)))
                                conn.commit(); imported+=1
                            st.success(f"Imported {imported} records for {team_name}"); st.rerun()
            except Exception as e: st.error(f"Error: {e}")
    # ── DELETE TEAM DATA ──
    with tab1:
        st.markdown('<div class="section-title" style="margin-top:32px;">Manage Uploaded Data</div>', unsafe_allow_html=True)
        cursor.execute("SELECT DISTINCT team FROM epl_players ORDER BY team")
        del_teams = [r[0] for r in cursor.fetchall()]
        if del_teams:
            st.markdown(f'<p style="font-size:12px;color:#6b7280;margin-bottom:12px;">{len(del_teams)} team(s) currently uploaded</p>', unsafe_allow_html=True)
            # Delete single team
            del_team = st.selectbox("Delete a specific team", ["-- Select team --"] + del_teams, key="del_team_sel")
            if del_team != "-- Select team --":
                st.warning(f"This permanently deletes all players and match data for **{del_team}**.")
                if st.button(f"Delete {del_team}", key="del_team_btn"):
                    pids_to_del = [r[0] for r in cursor.execute("SELECT id FROM epl_players WHERE team=?", (del_team,)).fetchall()]
                    for p in pids_to_del:
                        cursor.execute("DELETE FROM epl_sessions WHERE player_id=?", (p,))
                        cursor.execute("DELETE FROM epl_reports WHERE player_id=?", (p,))
                    cursor.execute("DELETE FROM epl_players WHERE team=?", (del_team,))
                    conn.commit()
                    st.success(f"{del_team} deleted.")
                    st.session_state.selected_epl_player_id = None
                    st.rerun()
            # Delete ALL data
            st.markdown('<hr style="border-top:1px solid #e5e7ef;margin:20px 0;">', unsafe_allow_html=True)
            st.markdown('<p style="font-size:12px;font-weight:600;color:#dc2626;margin-bottom:8px;">Danger Zone</p>', unsafe_allow_html=True)
            if st.button("Delete All Uploaded Pro Data", key="del_all_btn"):
                st.session_state["confirm_del_all"] = True
            if st.session_state.get("confirm_del_all"):
                st.error("This will permanently delete ALL teams, players, sessions and reports from Pro Data. This cannot be undone.")
                ca1, ca2, _ = st.columns([1, 1, 3])
                with ca1:
                    if st.button("Yes, delete everything", key="confirm_del_all_yes"):
                        cursor.execute("SELECT id FROM epl_players")
                        all_pids = [r[0] for r in cursor.fetchall()]
                        for p in all_pids:
                            cursor.execute("DELETE FROM epl_sessions WHERE player_id=?", (p,))
                            cursor.execute("DELETE FROM epl_reports WHERE player_id=?", (p,))
                        cursor.execute("DELETE FROM epl_players")
                        conn.commit()
                        st.session_state["confirm_del_all"] = False
                        st.session_state.selected_epl_player_id = None
                        st.success("All Pro Data deleted.")
                        st.rerun()
                with ca2:
                    if st.button("Cancel", key="cancel_del_all"):
                        st.session_state["confirm_del_all"] = False
                        st.rerun()
        else:
            st.caption("No professional data uploaded yet.")
    with tab2:
        if not st.session_state.selected_epl_player_id:
            st.markdown('<div class="no-select"><p style="font-size:40px;margin-bottom:12px">⚽</p><p>Select a player from the sidebar</p></div>', unsafe_allow_html=True)
        else:
            epid=st.session_state.selected_epl_player_id
            cursor.execute("SELECT * FROM epl_players WHERE id=?",(epid,)); ep=cursor.fetchone()
            cursor.execute("SELECT * FROM epl_sessions WHERE player_id=? ORDER BY gameweek",(epid,)); eps=cursor.fetchall()
            cursor.execute("SELECT * FROM epl_reports WHERE player_id=? ORDER BY created_at DESC LIMIT 1",(epid,)); epr=cursor.fetchone()
            st.markdown(f'<div class="page-header"><div><div class="page-title">{ep[1]}</div><div class="page-meta">{ep[3] or ""} &nbsp;·&nbsp; {ep[2]} &nbsp;·&nbsp; Age {ep[5] or "?"}</div></div></div>', unsafe_allow_html=True)
            if eps:
                tm=sum(sg(s,6) for s in eps); tg=sum(sg(s,7) for s in eps)
                ta=sum(sg(s,8) for s in eps); txg=round(sum(float(sg(s,11,0)) for s in eps),2)
                tsh=sum(sg(s,9) for s in eps); tp=sum(sg(s,13) for s in eps)
                c1,c2,c3,c4,c5,c6=st.columns(6)
                for col,lbl,val,style in [(c1,"Minutes",tm,"neutral"),(c2,"Goals",tg,"positive"),(c3,"Assists",ta,"positive"),(c4,"xG",txg,"highlight"),(c5,"Shots",tsh,"neutral"),(c6,"Passes",tp,"neutral")]:
                    with col: st.markdown(f'<div class="metric-card {style}"><div class="metric-label">{lbl}</div><div class="metric-value">{val}</div></div>', unsafe_allow_html=True)
                st.markdown('<div class="section-title">Match Data</div>', unsafe_allow_html=True)
                mrows=[{"GW":sg(s,2),"Opponent":sg(s,3,""),"Mins":sg(s,6),"Goals":sg(s,7),"Assists":sg(s,8),"Shots":sg(s,9),"xG":round(float(sg(s,11,0)),2),"Passes":sg(s,13),"Tackles":sg(s,21),"Rating":sg(s,26)} for s in eps]
                st.dataframe(pd.DataFrame(mrows),use_container_width=True,hide_index=True)
                st.markdown('<div class="section-title">Performance Charts</div>', unsafe_allow_html=True)
                gws=[f"GW{sg(s,2)}" for s in eps]
                cc1,cc2=st.columns(2)
                with cc1:
                    fig=go.Figure()
                    fig.add_trace(go.Bar(x=gws,y=[sg(s,7) for s in eps],name='Goals',marker_color='#1a3a8a',opacity=0.85,marker_line_width=0))
                    fig.add_trace(go.Bar(x=gws,y=[sg(s,8) for s in eps],name='Assists',marker_color='#3b82f6',opacity=0.85,marker_line_width=0))
                    fig.update_layout(title='Goals and Assists',barmode='group',paper_bgcolor='white',plot_bgcolor='white',margin=dict(l=10,r=10,t=40,b=10),xaxis=dict(showgrid=False,type='category'),yaxis=dict(gridcolor='#f3f4f6'),legend=dict(orientation='h',yanchor='bottom',y=1.02,xanchor='right',x=1),height=240)
                    st.plotly_chart(fig,use_container_width=True)
                with cc2:
                    fig=go.Figure()
                    fig.add_trace(go.Scatter(x=gws,y=[round(float(sg(s,11,0)),2) for s in eps],name='xG',line=dict(color='#f5c842',width=2.5),mode='lines+markers',marker=dict(size=7)))
                    fig.add_trace(go.Scatter(x=gws,y=[sg(s,7) for s in eps],name='Goals',line=dict(color='#1a3a8a',width=2,dash='dash'),mode='lines+markers',marker=dict(size=7)))
                    fig.update_layout(title='xG vs Actual Goals',paper_bgcolor='white',plot_bgcolor='white',margin=dict(l=10,r=10,t=40,b=10),xaxis=dict(showgrid=False,type='category'),yaxis=dict(gridcolor='#f3f4f6'),legend=dict(orientation='h',yanchor='bottom',y=1.02,xanchor='right',x=1),height=240)
                    st.plotly_chart(fig,use_container_width=True)
            st.markdown('<div class="section-title">Performance Audit Report</div>', unsafe_allow_html=True)
            if epr:
                render_report(epr[2],epid,ep[1],f"{ep[3] or ''} · {ep[2]}",is_pro=True)
            else:
                st.markdown('<div style="background:#f9fafb;border:1px solid #e5e7ef;border-radius:16px;padding:48px;text-align:center;"><p style="font-size:32px;margin-bottom:12px;">📊</p><p style="font-size:15px;font-weight:600;color:#374151;margin-bottom:24px;">No audit generated yet</p></div>', unsafe_allow_html=True)
                st.markdown('<div class="gen-btn" style="text-align:center;margin-top:16px;">', unsafe_allow_html=True)
                if st.button("Generate Performance Audit",key="gen_epl"):
                    if not eps: st.warning("No match data.")
                    else:
                        stext="\n".join([f"GW{sg(s,2)} vs {sg(s,3,'?')}: {sg(s,6)}mins, {sg(s,7)}G {sg(s,8)}A, {sg(s,9)} shots, xG {round(float(sg(s,11,0)),2)}, {sg(s,13)} passes, {sg(s,21)} tackles, rating {sg(s,26)}" for s in eps])
                        prompt=f"""You are the chief scout at a Premier League club. Professional performance audit for sporting director. Full sentences. No bullets. Every claim backed by data. Max 4 pages.

PLAYER: {ep[1]} | {ep[2]} | {ep[3] or 'Unknown'} | Age {ep[5] or '?'}
MATCH DATA ({len(eps)} gameweeks): {stext}

Sections: EXECUTIVE SUMMARY (4 sentences: quality+data, weakness+data, trajectory, transfer verdict) | 1. SEASON PERFORMANCE RATING | 2. GENERAL PERFORMANCE AUDIT | 3. TECHNICAL AND CREATIVE QUALITY | 4. DEFENSIVE CONTRIBUTION | 5. CONSISTENCY AND AVAILABILITY | 6. BEST POSITION AND TACTICAL FIT | 7. THREE STANDOUT PERFORMANCES | 8. SEASON TRAJECTORY | 9. INJURY AND FITNESS ASSESSMENT | 10. LEADERSHIP ASSESSMENT | 11. TRANSFER INTELLIGENCE (fee range, best fit club, negotiation note) | 12. CHIEF SCOUT VERDICT

Complete all sections. No truncation."""
                        with st.spinner("Generating audit..."):
                            rtext=ai_report(prompt)
                            cursor.execute("INSERT INTO epl_reports (player_id,report_text) VALUES (?,?)",(epid,rtext))
                            conn.commit(); st.rerun()
                st.markdown('</div>', unsafe_allow_html=True)
    with tab4:
        st.markdown('<div class="section-title">AI Smart Import</div>', unsafe_allow_html=True)
        st.markdown("""
        <div style="background:#eff6ff;border:1px solid #bfdbfe;border-radius:12px;padding:20px 24px;margin-bottom:24px;">
            <p style="font-size:14px;font-weight:700;color:#1d4ed8;margin-bottom:6px;">✨ Paste Any Professional Data</p>
            <p style="font-size:13px;color:#3b82f6;margin:0;line-height:1.7;">
            Paste data from any source — FBref, Transfermarkt, WhoScored, Sofascore, a PDF report, 
            a copied spreadsheet, or even handwritten notes. Claude will identify the players, 
            extract their stats, and structure them for the platform automatically.
            </p>
        </div>
        """, unsafe_allow_html=True)
        col_ex1, col_ex2 = st.columns(2)
        with col_ex1:
            st.markdown("""
            <div style="background:#f9fafb;border:1px solid #e5e7ef;border-radius:10px;padding:14px 16px;font-size:12px;color:#6b7280;">
            <strong style="color:#374151;">Works with:</strong><br>
            • Copied FBref tables<br>
            • Sofascore player stats<br>
            • PDF match reports (text)<br>
            • Spreadsheet pastes
            </div>
            """, unsafe_allow_html=True)
        with col_ex2:
            st.markdown("""
            <div style="background:#f9fafb;border:1px solid #e5e7ef;border-radius:10px;padding:14px 16px;font-size:12px;color:#6b7280;">
            <strong style="color:#374151;">Also works with:</strong><br>
            • WhatsApp scouting notes<br>
            • Rough text descriptions<br>
            • Mixed format data<br>
            • Any language
            </div>
            """, unsafe_allow_html=True)
        st.markdown('<div style="margin-top:20px;"></div>', unsafe_allow_html=True)
        pro_raw = st.text_area(
            "pro_raw",
            placeholder="""Paste anything here. Examples:

Player: Bruno Fernandes | Club: Man United | Pos: Midfielder | Age: 29
Season 2023/24: 35 apps, 1800 mins, 10 goals, 8 assists, 3 yellows

Or paste a full copied table from FBref or WhoScored...
Or type rough notes: Erling Haaland, 24, striker Man City, 27 goals 5 assists this season""",
            height=220,
            label_visibility="collapsed",
            key="pro_ai_raw"
        )
        if pro_raw and pro_raw.strip():
            st.markdown(f'<p style="font-size:11px;color:#9ca3af;">{len(pro_raw)} characters — ready to parse</p>', unsafe_allow_html=True)
            if st.button("✨ Parse and Import with AI", key="pro_ai_parse"):
                with st.spinner("Claude is reading your data..."):
                    try:
                        import json
                        result = ai_clean_data(pro_raw, "pro")
                        result = result.strip()
                        if result.startswith("```"):
                            result = result.split("```")[1]
                            if result.startswith("json"): result = result[4:]
                        result = result.strip()
                        parsed = json.loads(result)
                        if not isinstance(parsed, list): parsed = [parsed]
                        st.session_state["pro_ai_parsed"] = parsed
                        st.rerun()
                    except Exception as e:
                        st.error(f"Could not parse: {str(e)[:100]}")

        if st.session_state.get("pro_ai_parsed"):
            pro_players = st.session_state["pro_ai_parsed"]
            st.success(f"Claude detected {len(pro_players)} player(s) — confirm to save:")
            st.dataframe(pd.DataFrame([{"Name":p.get("name","?"),"Team":p.get("team","?"),"Position":p.get("position","?"),"Goals":p.get("goals",0),"Assists":p.get("assists",0)} for p in pro_players]), use_container_width=True, hide_index=True)
            sv1, sv2, _ = st.columns([1,1,4])
            with sv1:
                if st.button("✅ Save to Platform", key="pro_ai_confirm"):
                    import random; random.seed(99)
                    saved = 0; last_epid = None
                    for p in pro_players:
                        try:
                            pname = str(p.get("name","")).strip()
                            if not pname: continue
                            team = str(p.get("team","Unknown")).strip() or "Unknown"
                            appearances = max(int(p.get("appearances",1) or 1), 1)
                            total_goals = int(p.get("goals",0) or 0)
                            total_assists = int(p.get("assists",0) or 0)
                            total_mins = int(p.get("minutes", appearances*85) or appearances*85)
                            goals_p90 = float(p.get("goals_per_90",0) or 0)
                            assists_p90 = float(p.get("assists_per_90",0) or 0)
                            yellow_cards = int(p.get("yellow_cards",0) or 0)
                            xg_total = float(p.get("xg",0) or 0)
                            xa_total = float(p.get("xa",0) or 0)
                            cursor.execute("SELECT id FROM epl_players WHERE name=? AND team=?", (pname, team))
                            ex = cursor.fetchone()
                            if ex: epid = ex[0]
                            else:
                                cursor.execute("INSERT INTO epl_players (name,team,position,nationality,age) VALUES (?,?,?,?,?)",
                                    (pname, team, str(p.get("position","Midfielder")), str(p.get("nationality","")), int(p.get("age",0) or 0)))
                                conn.commit(); epid = cursor.lastrowid
                            last_epid = epid
                            goal_list = [0]*appearances; assist_list = [0]*appearances
                            for gi in range(total_goals): goal_list[random.randint(0,appearances-1)] += 1
                            for ai2 in range(total_assists): assist_list[random.randint(0,appearances-1)] += 1
                            avg_mins = int(total_mins/appearances)
                            xg_per = round(xg_total/appearances,2) if appearances>0 else 0
                            xa_per = round(xa_total/appearances,2) if appearances>0 else 0
                            for gw in range(1, appearances+1):
                                cursor.execute("INSERT INTO epl_sessions (player_id,gameweek,opponent,minutes_played,goals,assists,shots,shots_on_target,xg,xa,passes_completed,tackles_won,yellow_cards,red_cards,rating) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",
                                    (epid,gw,f"GW{gw}",max(45,min(90,avg_mins+random.randint(-5,5))),goal_list[gw-1],assist_list[gw-1],max(0,goal_list[gw-1]*3+random.randint(0,2)),max(0,goal_list[gw-1]*2),round(xg_per+random.uniform(-0.05,0.05),2),round(xa_per+random.uniform(-0.03,0.03),2),random.randint(25,55),random.randint(1,4),1 if gw<=yellow_cards else 0,0,round(6.5+(goals_p90+assists_p90)*1.5+random.uniform(-0.4,0.4),1)))
                            conn.commit(); saved += 1
                        except Exception as e:
                            st.warning(f"Skipped {p.get('name','?')}: {e}"); continue
                    st.session_state["pro_ai_parsed"] = None
                    if last_epid:
                        st.session_state.selected_epl_player_id = last_epid
                        st.session_state.mode = "pro"
                    st.success(f"✅ {saved} player(s) saved!"); st.rerun()
            with sv2:
                if st.button("✕ Discard", key="pro_ai_discard"):
                    st.session_state["pro_ai_parsed"] = None; st.rerun()
    with tab3:
        st.markdown('<div class="section-title">Team Analysis</div>', unsafe_allow_html=True)
        if not epl_teams:
            st.markdown('<div class="no-select"><p style="font-size:40px;margin-bottom:12px">⚽</p><p>Upload data in the Upload Data tab first</p></div>', unsafe_allow_html=True)
        else:
            sel_team=st.selectbox("Select team",epl_teams,key="st")
            if sel_team:
                cursor.execute("""SELECT p.name,p.position,p.age,SUM(s.minutes_played),SUM(s.goals),SUM(s.assists),ROUND(SUM(s.xg),2),SUM(s.shots),SUM(s.passes_completed),SUM(s.tackles_won),COUNT(s.id) FROM epl_players p LEFT JOIN epl_sessions s ON p.id=s.player_id WHERE p.team=? GROUP BY p.id ORDER BY SUM(s.goals) DESC""",(sel_team,))
                ts=cursor.fetchall()
                if ts:
                    cols=["Player","Position","Age","Minutes","Goals","Assists","xG","Shots","Passes","Tackles","Games"]
                    df_t=pd.DataFrame(ts,columns=cols)
                    st.dataframe(df_t,use_container_width=True,hide_index=True)
                    cc1,cc2=st.columns(2)
                    top8=df_t.nlargest(8,'Goals')
                    with cc1:
                        fig=go.Figure()
                        fig.add_trace(go.Bar(x=top8['Player'],y=top8['Goals'],marker_color='#1a3a8a',opacity=0.85,marker_line_width=0))
                        fig.update_layout(title='Goals by Player',paper_bgcolor='white',plot_bgcolor='white',margin=dict(l=10,r=10,t=40,b=80),xaxis=dict(showgrid=False,tickangle=45),yaxis=dict(gridcolor='#f3f4f6'),height=300,showlegend=False)
                        st.plotly_chart(fig,use_container_width=True)
                    with cc2:
                        fig=go.Figure()
                        fig.add_trace(go.Bar(x=top8['Player'],y=top8['xG'],name='xG',marker_color='#f5c842',opacity=0.85,marker_line_width=0))
                        fig.add_trace(go.Bar(x=top8['Player'],y=top8['Goals'],name='Goals',marker_color='#1a3a8a',opacity=0.7,marker_line_width=0))
                        fig.update_layout(title='xG vs Goals',barmode='group',paper_bgcolor='white',plot_bgcolor='white',margin=dict(l=10,r=10,t=40,b=80),xaxis=dict(showgrid=False,tickangle=45),yaxis=dict(gridcolor='#f3f4f6'),height=300,legend=dict(orientation='h',yanchor='bottom',y=1.02,xanchor='right',x=1))
                        st.plotly_chart(fig,use_container_width=True)
